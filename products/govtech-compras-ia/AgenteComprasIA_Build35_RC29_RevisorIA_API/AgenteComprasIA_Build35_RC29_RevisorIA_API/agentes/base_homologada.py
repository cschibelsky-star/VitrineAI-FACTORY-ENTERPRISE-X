from pathlib import Path
from datetime import datetime
import json
import re
import shutil
import unicodedata

try:
    from docx import Document
except Exception:
    Document = None


ROOT = Path(__file__).resolve().parents[1]
BASE_DIR = ROOT / "base_conhecimento" / "homologada"
CASOS_FILE = BASE_DIR / "casos_homologados.jsonl"
DOCS_DIR = BASE_DIR / "documentos"


def _ensure_dirs():
    BASE_DIR.mkdir(parents=True, exist_ok=True)
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    if not CASOS_FILE.exists():
        CASOS_FILE.write_text("", encoding="utf-8")


def normalizar(txt):
    txt = str(txt or "").lower().strip()
    txt = unicodedata.normalize("NFKD", txt)
    txt = "".join(c for c in txt if not unicodedata.combining(c))
    txt = re.sub(r"[^a-z0-9]+", " ", txt)
    return re.sub(r"\s+", " ", txt).strip()


def slug(txt, limite=80):
    txt = normalizar(txt)
    txt = re.sub(r"\s+", "_", txt).strip("_")
    return (txt[:limite].strip("_") or "caso")


def ler_docx_texto(path):
    if not Document:
        return ""
    try:
        doc = Document(str(path))
        partes = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        partes.append(cell.text)
        return "\n".join(partes)
    except Exception:
        return ""


def carregar_casos():
    _ensure_dirs()
    casos = []
    for linha in CASOS_FILE.read_text(encoding="utf-8", errors="ignore").splitlines():
        linha = linha.strip()
        if not linha:
            continue
        try:
            casos.append(json.loads(linha))
        except Exception:
            continue
    return casos


def salvar_caso(caso):
    _ensure_dirs()
    caso = dict(caso)
    caso.setdefault("id", datetime.now().strftime("%Y%m%d%H%M%S"))
    caso.setdefault("data_homologacao", datetime.now().isoformat())
    with CASOS_FILE.open("a", encoding="utf-8") as f:
        f.write(json.dumps(caso, ensure_ascii=False) + "\n")
    return caso


def detectar_tipo_contratacao(dados, classificacao=None):
    texto = " ".join([
        str((dados or {}).get("descricao", "")),
        str((dados or {}).get("artista", "")),
        str((dados or {}).get("evento", "")),
        str((classificacao or {}).get("modalidade", "")),
        str((classificacao or {}).get("categoria", "")),
    ])
    low = normalizar(texto)
    if "inexigibilidade" in low or "artistic" in low or "cantor" in low or "cantora" in low or "show" in low or "apresentacao" in low:
        return "contratacao_artistica_inexigibilidade"
    if "dispensa" in low:
        return "dispensa"
    return "geral"


def extrair_resumo_documentos(pasta_processo):
    pasta = Path(pasta_processo)
    resumo = {}
    if not pasta.exists():
        return resumo

    for docx in list(pasta.rglob("*.docx")):
        nome = docx.name
        if nome.startswith("~$"):
            continue
        texto = ler_docx_texto(docx)
        if texto:
            resumo[nome] = {
                "arquivo": str(docx),
                "texto_amostra": texto[:5000],
            }
    return resumo


def adicionar_processo_homologado(pasta_processo, dados, classificacao=None, observacoes=""):
    """
    Guarda o processo como caso homologado.
    Importante: não treina IA nem altera automaticamente regras jurídicas.
    Apenas cria referência estruturada para consulta nos próximos processos.
    """
    _ensure_dirs()
    dados = dict(dados or {})
    classificacao = dict(classificacao or {})
    tipo = detectar_tipo_contratacao(dados, classificacao)

    nome_ref = dados.get("artista") or dados.get("fornecedor") or dados.get("evento") or dados.get("descricao") or "processo"
    caso_id = datetime.now().strftime("%Y%m%d%H%M%S") + "_" + slug(nome_ref, 40)
    destino = DOCS_DIR / caso_id
    destino.mkdir(parents=True, exist_ok=True)

    pasta_processo = Path(pasta_processo)
    arquivos_copiados = []
    if pasta_processo.exists():
        for arq in pasta_processo.rglob("*"):
            if arq.is_file() and arq.suffix.lower() in [".docx", ".pdf", ".txt"]:
                rel = arq.relative_to(pasta_processo)
                dest = destino / rel
                dest.parent.mkdir(parents=True, exist_ok=True)
                try:
                    shutil.copy2(arq, dest)
                    arquivos_copiados.append(str(dest.relative_to(ROOT)))
                except Exception:
                    pass

    caso = {
        "id": caso_id,
        "tipo": tipo,
        "modalidade": classificacao.get("modalidade", ""),
        "categoria": classificacao.get("categoria", ""),
        "artista": dados.get("artista", ""),
        "evento": dados.get("evento", ""),
        "objeto": dados.get("descricao", ""),
        "valor": dados.get("valor_estimado", ""),
        "secretaria": dados.get("secretaria", ""),
        "processo": dados.get("numero_processo", ""),
        "requisicao": dados.get("numero_requisicao", ""),
        "data_documento": dados.get("data_documento", ""),
        "data_evento": dados.get("data_evento", ""),
        "observacoes": observacoes,
        "arquivos": arquivos_copiados,
        "resumo_documentos": extrair_resumo_documentos(pasta_processo),
    }
    return salvar_caso(caso)


def pontuar_similaridade(caso, dados, classificacao=None):
    dados = dados or {}
    classificacao = classificacao or {}
    score = 0

    if caso.get("tipo") == detectar_tipo_contratacao(dados, classificacao):
        score += 50
    if normalizar(caso.get("modalidade")) and normalizar(caso.get("modalidade")) == normalizar(classificacao.get("modalidade")):
        score += 15
    if normalizar(caso.get("categoria")) and normalizar(caso.get("categoria")) == normalizar(classificacao.get("categoria")):
        score += 10
    if normalizar(caso.get("secretaria")) and normalizar(caso.get("secretaria")) == normalizar(dados.get("secretaria")):
        score += 10
    if caso.get("artista") and normalizar(caso.get("artista")) == normalizar(dados.get("artista")):
        score += 20
    if normalizar(caso.get("evento")) and normalizar(caso.get("evento")) == normalizar(dados.get("evento")):
        score += 10

    texto_caso = normalizar(" ".join([caso.get("objeto", ""), caso.get("evento", ""), caso.get("artista", "")]))
    texto_atual = normalizar(" ".join([dados.get("descricao", ""), dados.get("evento", ""), dados.get("artista", "")]))
    comuns = set(texto_caso.split()) & set(texto_atual.split())
    score += min(len(comuns), 10)

    return score


def buscar_caso_similar(dados, classificacao=None):
    casos = carregar_casos()
    if not casos:
        return None

    ranqueados = []
    for c in casos:
        s = pontuar_similaridade(c, dados, classificacao)
        if s > 0:
            ranqueados.append((s, c))
    if not ranqueados:
        return None
    ranqueados.sort(key=lambda x: x[0], reverse=True)
    score, caso = ranqueados[0]
    if score < 45:
        return None
    caso = dict(caso)
    caso["score_similaridade"] = score
    return caso


def resumo_caso(caso):
    if not caso:
        return "Nenhum caso homologado semelhante localizado."
    partes = [
        f"Referência homologada: {caso.get('artista') or caso.get('evento') or caso.get('id')}",
        f"Tipo: {caso.get('tipo', '')}",
        f"Modalidade: {caso.get('modalidade', '')}",
        f"Valor: {caso.get('valor', '')}",
        f"Processo: {caso.get('processo', '')}",
        f"Similaridade: {caso.get('score_similaridade', '')}",
    ]
    return " | ".join([p for p in partes if p and not p.endswith(": ")])