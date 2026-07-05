from pathlib import Path
from datetime import datetime
import json
import os
import re
import urllib.request
import urllib.error

from agentes.base_conhecimento import gerar_contexto_base_para_revisor

try:
    from docx import Document
except Exception:
    Document = None


MAX_CHARS_DOC = 18000


def ler_docx(path):
    if Document is None:
        return ""
    try:
        doc = Document(str(path))
        partes = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                partes.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                linha = []
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        linha.append(cell.text.strip())
                if linha:
                    partes.append(" | ".join(linha))
        return "\n".join(partes)
    except Exception as e:
        return f"[ERRO AO LER DOCX: {e}]"


def salvar_docx_texto(path, titulo, texto):
    if Document is None:
        Path(path).write_text(texto or "", encoding="utf-8")
        return Path(path)

    doc = Document()
    doc.add_heading(titulo, level=1)
    for bloco in (texto or "").split("\n"):
        bloco = bloco.rstrip()
        if not bloco:
            doc.add_paragraph("")
        elif bloco.startswith("# "):
            doc.add_heading(bloco[2:].strip(), level=1)
        elif bloco.startswith("## "):
            doc.add_heading(bloco[3:].strip(), level=2)
        elif bloco.startswith("### "):
            doc.add_heading(bloco[4:].strip(), level=3)
        else:
            doc.add_paragraph(bloco)
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def documentos_principais(pasta_processo):
    pasta = Path(pasta_processo)
    candidatos = []
    for nome in ["02_DFD.docx", "03_ETP.docx", "04_TR.docx"]:
        achados = list(pasta.rglob(nome))
        if achados:
            candidatos.append(achados[0])
    return candidatos


def montar_pacote_texto(pasta_processo):
    docs = documentos_principais(pasta_processo)
    partes = []
    for d in docs:
        txt = ler_docx(d)
        if len(txt) > MAX_CHARS_DOC:
            txt = txt[:MAX_CHARS_DOC] + "\n[DOCUMENTO TRUNCADO PARA REVISÃO]"
        partes.append(f"\n\n===== {d.name} =====\n{txt}")
    return "\n".join(partes), docs


def mascarar_dados_sensiveis(texto):
    if not texto:
        return ""
    s = texto
    s = re.sub(r"\b\d{3}\.\d{3}\.\d{3}-\d{2}\b", "[CPF MASCARADO]", s)
    s = re.sub(r"\b\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}\b", "[CNPJ MASCARADO]", s)
    s = re.sub(r"\b\d{2}\.\d{3}\.\d{3}-\d\b", "[RG MASCARADO]", s)
    return s


def prompt_revisao(pacote_texto, dados=None, contexto_base=''):
    dados = dados or {}
    return f"""
Você é um revisor técnico-jurídico especializado em fase preparatória de contratação pública municipal,
com foco em contratação artística por inexigibilidade de licitação e processos de adesão à Ata de Registro de Preços.

TAREFA:
Revise os documentos DFD, ETP e TR abaixo.

REGRAS:
- Não invente dados.
- Não altere valores, datas, processo, requisição, dotação ou nomes se não houver inconsistência clara.
- Se faltar dado, marque como pendência.
- Para contratação artística, conferir aderência ao art. 74, inciso II, da Lei 14.133/2021.
- Para adesão à Ata de Registro de Preços, conferir vantajosidade, compatibilidade do objeto, vigência da Ata, saldo, anuência do órgão gerenciador, aceite da empresa, regularidade da empresa, dotação, parecer jurídico, autorização, empenho e publicação do contrato/extrato quando aplicável.
- Não usar "notória especialização" para contratação artística.
- Apontar erros gramaticais, concordância de gênero, duplicidade textual, placeholder e frases mal estruturadas.
- Conferir coerência entre DFD, ETP e TR.
- Separar achados em: CRÍTICO, ATENÇÃO e SUGESTÃO.
- Ao final, informar se recomenda ou não homologar na base de conhecimento.

CONTEXTO DA BASE DE CONHECIMENTO LOCAL:
{contexto_base or "Nenhuma referência local adicional localizada."}

DADOS INFORMADOS PELO SISTEMA:
Artista: {dados.get("artista", "")}
Evento: {dados.get("evento", "")}
Valor: {dados.get("valor_estimado", "")}
Processo: {dados.get("numero_processo", "")}
Requisição: {dados.get("numero_requisicao", "")}
Secretaria: {dados.get("secretaria", "")}
Data do evento: {dados.get("data_evento", "")}
Data do documento: {dados.get("data_documento", "")}

DOCUMENTOS:
{pacote_texto}

FORMATO DA RESPOSTA:
# RELATÓRIO DE REVISÃO IA

## Resultado geral
Apto / Apto com ressalvas / Não apto

## Críticos
- ...

## Atenções
- ...

## Sugestões
- ...

## Melhorias recomendadas por documento
### DFD
- ...
### ETP
- ...
### TR
- ...

## Recomendação para homologação
Homologar / Não homologar / Homologar somente após correções

## Checklist final
- Gramática:
- Datas:
- Valores:
- Artista:
- Modalidade:
- Fundamentação:
- Campos administrativos:
"""


def prompt_melhoria_documento(nome_doc, texto_doc, relatorio, dados=None):
    dados = dados or {}
    return f"""
Você é um redator técnico-jurídico de contratações públicas municipais.
Melhore o texto do documento {nome_doc}, preservando a estrutura oficial, títulos, numeração e dados variáveis.

REGRAS OBRIGATÓRIAS:
- Não invente número de processo, valor, data, dotação, artista, local ou autoridade.
- Preserve a finalidade pública e o padrão de controle interno/externo.
- Corrija somente gramática, clareza, coesão, estrutura textual e inconsistências apontadas.
- Não usar "notória especialização" em contratação artística.
- Em adesão à Ata, não tratar como contratação direta artística nem como dispensa comum.
- Se o artista for feminino, usar cantora/da cantora quando aplicável.
- Se o artista for masculino, usar cantor/do cantor quando aplicável.
- Se o artista não estiver identificado, usar "do(a) artista indicado(a)", sem duplicidade.
- Não criar novas obrigações que não estejam na lógica do processo.
- A resposta deve ser o texto revisado em formato simples, sem comentários externos.

DADOS:
Artista: {dados.get("artista", "")}
Evento: {dados.get("evento", "")}
Valor: {dados.get("valor_estimado", "")}
Processo: {dados.get("numero_processo", "")}
Requisição: {dados.get("numero_requisicao", "")}
Secretaria: {dados.get("secretaria", "")}

RELATÓRIO DE REVISÃO:
{relatorio[:6000]}

DOCUMENTO ORIGINAL:
{texto_doc[:MAX_CHARS_DOC]}

RETORNE APENAS O TEXTO REVISADO:
"""


def chamar_api_openai_compat(prompt, api_key, endpoint, model, temperature=0.1, timeout=90):
    endpoint = endpoint or "https://api.openai.com/v1/chat/completions"
    model = model or "gpt-4o-mini"
    payload = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "Você é um revisor técnico-jurídico rigoroso, objetivo e conservador. Não invente dados."
            },
            {"role": "user", "content": prompt},
        ],
        "temperature": temperature,
    }
    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(
        endpoint,
        data=data,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            raw = resp.read().decode("utf-8", errors="ignore")
            obj = json.loads(raw)
            return obj["choices"][0]["message"]["content"], None
    except urllib.error.HTTPError as e:
        try:
            detalhe = e.read().decode("utf-8", errors="ignore")
        except Exception:
            detalhe = str(e)
        return "", f"Erro HTTP da API: {e.code} - {detalhe}"
    except Exception as e:
        return "", f"Erro ao chamar API: {e}"


def revisar_processo_com_ia(pasta_processo, dados, api_key, endpoint=None, model=None, mascarar=True):
    pacote, docs = montar_pacote_texto(pasta_processo)
    if not docs:
        return {
            "ok": False,
            "erro": "Não encontrei DFD, ETP e TR na pasta do processo.",
            "arquivos": [],
        }

    pacote_envio = mascarar_dados_sensiveis(pacote) if mascarar else pacote
    contexto_base = gerar_contexto_base_para_revisor(dados, {})
    contexto_base_envio = mascarar_dados_sensiveis(contexto_base) if mascarar else contexto_base
    prompt = prompt_revisao(pacote_envio, dados, contexto_base_envio)
    relatorio, erro = chamar_api_openai_compat(prompt, api_key, endpoint, model)
    if erro:
        return {"ok": False, "erro": erro, "arquivos": []}

    pasta = Path(pasta_processo)
    out_dir = pasta / "REVISAO_IA"
    out_dir.mkdir(parents=True, exist_ok=True)

    arqs = []
    rel_doc = salvar_docx_texto(out_dir / "Relatorio_Revisao_IA.docx", "Relatório de Revisão IA", relatorio)
    arqs.append(rel_doc)

    # Gera versões textuais revisadas para revisão humana.
    for docx in docs:
        txt = ler_docx(docx)
        txt_envio = mascarar_dados_sensiveis(txt) if mascarar else txt
        p2 = prompt_melhoria_documento(docx.name, txt_envio, relatorio, dados)
        revisado, erro2 = chamar_api_openai_compat(p2, api_key, endpoint, model)
        if erro2:
            revisado = f"Não foi possível gerar versão revisada de {docx.name}.\n\nErro: {erro2}\n\nDocumento original mantido para revisão manual."
        nome_saida = docx.stem + "_Revisado_IA.docx"
        arq = salvar_docx_texto(out_dir / nome_saida, f"{docx.stem} — Versão Revisada por IA", revisado)
        arqs.append(arq)

    manifest = {
        "data_revisao": datetime.now().isoformat(),
        "pasta_processo": str(pasta_processo),
        "modelo": model or "gpt-4o-mini",
        "endpoint": endpoint or "https://api.openai.com/v1/chat/completions",
        "mascaramento": bool(mascarar),
        "base_conhecimento_consultada": True,
        "arquivos": [str(a) for a in arqs],
    }
    (out_dir / "manifesto_revisao_ia.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    arqs.append(out_dir / "manifesto_revisao_ia.json")

    return {
        "ok": True,
        "erro": "",
        "relatorio": relatorio,
        "arquivos": arqs,
        "pasta_revisao": out_dir,
    }