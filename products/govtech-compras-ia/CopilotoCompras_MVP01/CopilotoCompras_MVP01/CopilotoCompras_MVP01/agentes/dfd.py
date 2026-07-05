from pathlib import Path
from datetime import datetime
from docx import Document


def gerar_dfd(processo_id: int, descricao: str, classificacao: dict, valor_estimado: str, pasta_saida: Path) -> Path:
    pasta_saida.mkdir(parents=True, exist_ok=True)
    arquivo = pasta_saida / "DFD_inicial.docx"

    doc = Document()
    doc.add_heading("DOCUMENTO DE FORMALIZAÇÃO DE DEMANDA", level=1)

    doc.add_paragraph("SECRETARIA MUNICIPAL DE SUMARÉ")
    doc.add_paragraph(f"Processo local nº: {processo_id}")
    doc.add_paragraph("")

    doc.add_heading("1. IDENTIFICAÇÃO DA DEMANDA", level=2)
    doc.add_paragraph(f"1.1. Tipo: SERVIÇO" if "serviço" in classificacao.get("categoria", "").lower() else "1.1. Tipo: INFORMAÇÃO PENDENTE")
    doc.add_paragraph(f"1.2. Classe: {classificacao.get('categoria', 'INFORMAÇÃO PENDENTE')}")
    doc.add_paragraph(f"1.3. Objeto: {descricao}")

    doc.add_heading("2. DESCRIÇÃO SUCINTA DO OBJETO", level=2)
    doc.add_paragraph(classificacao.get("objeto_resumido", descricao))

    doc.add_heading("3. QUANTIDADE A SER CONTRATADA", level=2)
    doc.add_paragraph("01 unidade/apresentação/serviço, conforme detalhamento posterior no Termo de Referência.")

    doc.add_heading("4. ESTIMATIVA PRELIMINAR DO VALOR DA CONTRATAÇÃO", level=2)
    doc.add_paragraph(valor_estimado if valor_estimado else "INFORMAÇÃO PENDENTE")

    doc.add_heading("5. JUSTIFICATIVA DA NECESSIDADE DA CONTRATAÇÃO", level=2)
    if "artística" in descricao.lower() or "show" in descricao.lower() or "cantor" in descricao.lower() or "cantora" in descricao.lower():
        justificativa = (
            "A presente contratação tem por finalidade atender às demandas da Secretaria Municipal responsável "
            "pela realização de evento institucional, cultural ou comunitário, com relevante papel na promoção "
            "da cultura, lazer, integração social e fortalecimento das atividades públicas do Município. "
            "A contratação deverá ser instruída com os documentos comprobatórios necessários, incluindo proposta, "
            "justificativa de preço e documentos que demonstrem a regularidade e a adequação da escolha."
        )
    else:
        justificativa = (
            "A presente demanda visa atender necessidade administrativa da secretaria requisitante, "
            "devendo ser complementada com justificativa técnica específica, estimativa de valor, "
            "quantidade e demais elementos necessários à adequada instrução processual."
        )
    doc.add_paragraph(justificativa)

    doc.add_heading("6. VINCULAÇÃO OU DEPENDÊNCIA COM OUTRO DFD", level=2)
    doc.add_paragraph("Não identificada nesta fase inicial. Confirmar na revisão.")

    doc.add_heading("7. DATA PRETENDIDA PARA CONCLUSÃO DA CONTRATAÇÃO", level=2)
    doc.add_paragraph("INFORMAÇÃO PENDENTE")

    doc.add_heading("8. GRAU DE PRIORIDADE", level=2)
    doc.add_paragraph("Alta, média ou baixa: INFORMAÇÃO PENDENTE")

    doc.add_heading("9. ÁREA REQUISITANTE", level=2)
    doc.add_paragraph(classificacao.get("secretaria", "INFORMAÇÃO PENDENTE"))

    doc.add_heading("10. PENDÊNCIAS IDENTIFICADAS PELO COPILOTO", level=2)
    pendencias = classificacao.get("pendencias", [])
    if pendencias:
        for p in pendencias:
            doc.add_paragraph(f"- {p}")
    else:
        doc.add_paragraph("Nenhuma pendência inicial identificada.")

    doc.add_paragraph("")
    doc.add_paragraph(f"Sumaré, {datetime.now().strftime('%d/%m/%Y')}.")
    doc.add_paragraph("")
    doc.add_paragraph("__________________________________")
    doc.add_paragraph("Responsável pela demanda")

    doc.save(arquivo)
    return arquivo