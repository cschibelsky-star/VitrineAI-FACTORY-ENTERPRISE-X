from pathlib import Path
from datetime import datetime
from docx import Document


def gerar_etp(processo_id: int, descricao: str, classificacao: dict, valor_estimado: str, pasta_saida: Path) -> Path:
    pasta_saida.mkdir(parents=True, exist_ok=True)
    arquivo = pasta_saida / "ETP_inicial.docx"

    doc = Document()
    doc.add_heading("ESTUDO TÉCNICO PRELIMINAR", level=1)
    doc.add_paragraph("SECRETARIA MUNICIPAL DE SUMARÉ")
    doc.add_paragraph(f"Processo local nº: {processo_id}")

    doc.add_heading("1. DESCRIÇÃO DA NECESSIDADE DA CONTRATAÇÃO", level=2)
    doc.add_paragraph(
        f"A presente demanda refere-se à {descricao}. "
        "A contratação visa atender necessidade pública da secretaria requisitante, "
        "observando o planejamento administrativo, a continuidade dos serviços públicos "
        "e o interesse público envolvido."
    )

    doc.add_heading("2. PREVISÃO DA CONTRATAÇÃO NO PLANO DE CONTRATAÇÕES ANUAL", level=2)
    doc.add_paragraph(
        "INFORMAÇÃO PENDENTE. Confirmar se a contratação consta no Plano de Contratações Anual vigente."
    )

    doc.add_heading("3. REQUISITOS DA CONTRATAÇÃO", level=2)
    doc.add_paragraph(
        "Os requisitos da contratação deverão ser detalhados no Termo de Referência, "
        "incluindo condições de execução, obrigações da contratada, habilitação e demais exigências aplicáveis."
    )

    doc.add_heading("4. ESTIMATIVA DA QUANTIDADE", level=2)
    doc.add_paragraph(
        "01 unidade/apresentação/serviço, conforme detalhamento posterior no Termo de Referência."
    )

    doc.add_heading("5. LEVANTAMENTO DE MERCADO", level=2)
    if "inexigibilidade" in classificacao.get("modalidade", "").lower():
        doc.add_paragraph(
            "Tratando-se de possível contratação por inexigibilidade, o levantamento de mercado deverá "
            "priorizar a comprovação da compatibilidade do preço com contratações similares, notas fiscais, "
            "contratos anteriores, propostas e demais documentos capazes de demonstrar a razoabilidade do valor."
        )
    else:
        doc.add_paragraph(
            "O levantamento de mercado deverá considerar pesquisa de preços, contratações similares, "
            "consultas a fornecedores e demais parâmetros admitidos pela legislação aplicável."
        )

    doc.add_heading("6. ESTIMATIVA DO VALOR DA CONTRATAÇÃO", level=2)
    doc.add_paragraph(valor_estimado if valor_estimado else "INFORMAÇÃO PENDENTE")

    doc.add_heading("7. DESCRIÇÃO DA SOLUÇÃO COMO UM TODO", level=2)
    doc.add_paragraph(
        classificacao.get("objeto_resumido", descricao)
    )

    doc.add_heading("8. JUSTIFICATIVA PARA O PARCELAMENTO OU NÃO DA CONTRATAÇÃO", level=2)
    doc.add_paragraph(
        "A contratação deverá ser avaliada quanto à possibilidade de parcelamento. "
        "No caso de apresentação artística ou obrigação de resultado único, o parcelamento tende a não ser aplicável, "
        "por se tratar de execução indivisível vinculada ao objeto principal."
    )

    doc.add_heading("9. RESULTADOS PRETENDIDOS", level=2)
    doc.add_paragraph(
        "Pretende-se atender ao interesse público, assegurar a adequada execução do objeto, "
        "promover eficiência administrativa e garantir resultado compatível com a necessidade apresentada."
    )

    doc.add_heading("10. PROVIDÊNCIAS PRÉVIAS DA ADMINISTRAÇÃO", level=2)
    doc.add_paragraph(
        "Deverão ser adotadas as providências necessárias à instrução processual, incluindo "
        "reserva orçamentária, designação de gestor e fiscal, elaboração do Termo de Referência, "
        "validação da documentação e aprovação da autoridade competente."
    )

    doc.add_heading("11. CONTRATAÇÕES CORRELATAS OU INTERDEPENDENTES", level=2)
    doc.add_paragraph(
        "INFORMAÇÃO PENDENTE. Avaliar se existem contratações correlatas, como estrutura, logística, transporte, "
        "segurança, divulgação, limpeza, equipamentos ou outros serviços relacionados."
    )

    doc.add_heading("12. IMPACTOS AMBIENTAIS E MEDIDAS MITIGADORAS", level=2)
    doc.add_paragraph(
        "Os impactos ambientais deverão ser avaliados conforme a natureza do objeto. "
        "Quando aplicável, recomenda-se adoção de medidas de redução de resíduos, uso racional de recursos "
        "e destinação adequada de materiais."
    )

    doc.add_heading("13. VIABILIDADE DA CONTRATAÇÃO", level=2)
    doc.add_paragraph(
        "Diante das informações preliminares, a contratação mostra-se potencialmente viável, "
        "condicionada à complementação das informações pendentes, validação documental, disponibilidade orçamentária "
        "e análise jurídica competente."
    )

    doc.add_heading("14. PENDÊNCIAS IDENTIFICADAS PELO COPILOTO", level=2)
    pendencias = classificacao.get("pendencias", [])
    if pendencias:
        for p in pendencias:
            doc.add_paragraph(f"- {p}")
    else:
        doc.add_paragraph("Nenhuma pendência inicial identificada.")

    doc.add_paragraph("")
    doc.add_paragraph(f"Sumaré, {datetime.now().strftime('%d/%m/%Y')}.")
    doc.add_paragraph("")
    doc.add_paragraph("__________________________________")
    doc.add_paragraph("Agente responsável pela elaboração do ETP")

    doc.save(arquivo)
    return arquivo