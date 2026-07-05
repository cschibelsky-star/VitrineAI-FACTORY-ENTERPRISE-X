from pathlib import Path
from datetime import datetime
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
MODELOS = ROOT / "modelos_prefeitura"
COMPLEMENTAR = "A informar pela unidade responsável"

# Padrão definido para os processos da Secretaria Municipal de Cultura e Turismo.
GESTOR_PADRAO = "Carla Andressa Dourado"
GESTOR_CARGO_PADRAO = "Diretora de Divisão"
GESTOR_COMPETENCIAS_PADRAO = "Vasto conhecimento sobre as rotinas dos eventos da Secretaria Municipal de Cultura"

FISCAL_PADRAO = "Talita Cristiane Carvalho"
FISCAL_CARGO_PADRAO = "Diretora de Área"
FISCAL_COMPETENCIAS_PADRAO = "Vasto conhecimento sobre as rotinas dos eventos da Secretaria Municipal de Cultura"

# Cadastro interno de agentes conhecidos. A Requisição continua sendo a fonte do nome,
# mas matrícula e cargo podem ser completados por este mapa.
AGENTES_CONHECIDOS = {
    "CRISTIAN MARCELO SCHIBELSKY": {
        "nome": "Cristian Marcelo Schibelsky",
        "matricula": "22019",
        "cargo": "Diretor de Sub Divisão",
    },
    "CRISTIAN MARCELO SCHIBELSKY MATIOLI": {
        "nome": "Cristian Marcelo Schibelsky",
        "matricula": "22019",
        "cargo": "Diretor de Sub Divisão",
    },
}





# =========================
# FORMATAÇÃO / LAYOUT
# =========================

FONTE_PADRAO = "Arial"
TAM_BODY = 11
TAM_TABELA = 10


def cm(x):
    return Inches(x / 2.54)


def set_cell_width(cell, width):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.first_child_found_in("w:tcW")
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width.inches * 1440)))
    tcW.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill="D9EAF7"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text_format(cell, bold=False, font_size=TAM_TABELA):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1
        for run in p.runs:
            run.font.name = FONTE_PADRAO
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONTE_PADRAO)
            run.font.size = Pt(font_size)
            run.bold = bold or run.bold


def formatar_paragrafo(p, tipo="body"):
    pf = p.paragraph_format
    pf.line_spacing = 1.08
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    if tipo == "titulo":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(0)
        pf.space_after = Pt(8)
        pf.keep_with_next = True
        size = 12
        bold = True
    elif tipo == "subtitulo":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_after = Pt(4)
        pf.keep_with_next = True
        size = 11
        bold = True
    elif tipo == "secao":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(10)
        pf.space_after = Pt(4)
        pf.keep_with_next = True
        size = 11
        bold = True
    elif tipo == "item":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.left_indent = cm(0.6)
        pf.first_line_indent = cm(-0.25)
        pf.space_after = Pt(2)
        size = 11
        bold = False
    elif tipo == "assinatura":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        size = 11
        bold = False
    elif tipo == "campo":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_after = Pt(3)
        size = 11
        bold = False
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        size = 11
        bold = False
    for r in p.runs:
        r.font.name = FONTE_PADRAO
        r._element.rPr.rFonts.set(qn('w:eastAsia'), FONTE_PADRAO)
        r.font.size = Pt(size)
        if bold:
            r.bold = True


def identificar_tipo_paragrafo(texto, bold_hint=False):
    t = (texto or "").strip()
    if not t:
        return "campo"
    if t in {"DOCUMENTO DE FORMALIZAÇÃO DE DEMANDA", "ESTUDO TÉCNICO PRELIMINAR", "TERMO DE REFERÊNCIA"}:
        return "titulo"
    if t in {"Serviços sem dedicação exclusiva de mão-de-obra", "Licitação ou Contratação Direta"}:
        return "subtitulo"
    if t.startswith("TERMO DE APROVAÇÃO") or t.startswith("DISPOSIÇÕES FINAIS") or t.startswith("REFERÊNCIAS"):
        return "secao"
    if t.startswith("· "):
        return "item"
    if t.startswith("_") or t.startswith("Assinatura") or t.startswith("Nome:") or t.startswith("Matrícula:") or t.startswith("Mat.") or t.startswith("Sumaré,"):
        return "assinatura"
    if bold_hint or (len(t) > 3 and (t[0].isdigit() or t.startswith("Exigências")) and (" – " in t or " - " in t or t.isupper() or t.startswith("Exigências"))):
        return "secao"
    return "body"


# Metadados de rodapé dos modelos oficiais atualmente adotados.
# Quando a Prefeitura atualizar o modelo, alterar aqui uma única vez.
RODAPE_MODELOS_ATUAIS = {
    "DFD": {
        "titulo": "Documento de Formalização de Demanda",
        "atualizacao": "",
    },
    "ETP": {
        "titulo": "Estudo Técnico Preliminar",
        "atualizacao": "atualização: novembro/2023 v.01",
    },
    "TR": {
        "titulo": "Termo de Referência – Serviços sem dedicação exclusiva de mão de obra",
        "atualizacao": "Atualização: AGO/25",
    },
}


def add_field_run(paragraph, instr):
    """Insere campo Word/LibreOffice, usado para PAGE e NUMPAGES."""
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    run_instr = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " " + instr + " "
    run_instr._r.append(instr_text)

    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)

    run_cache = paragraph.add_run("1")

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)

    return run_cache


def aplicar_fonte_rodape(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    for r in paragraph.runs:
        r.font.name = FONTE_PADRAO
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONTE_PADRAO)
        r.font.size = Pt(8)


def limpar_footer(footer):
    # Remove parágrafos/tabelas herdados do modelo, inclusive paginação estática "Página 11 | 11".
    for child in list(footer._element):
        footer._element.remove(child)


def aplicar_rodape_modelo_atual(doc, tipo_documento="GERAL"):
    meta = RODAPE_MODELOS_ATUAIS.get(tipo_documento)
    if not meta:
        return

    for section in doc.sections:
        footer = section.footer
        limpar_footer(footer)

        p_pag = footer.add_paragraph()
        p_pag.add_run("Página ")
        add_field_run(p_pag, "PAGE")
        p_pag.add_run(" | ")
        add_field_run(p_pag, "NUMPAGES")
        aplicar_fonte_rodape(p_pag)

        p_titulo = footer.add_paragraph(meta["titulo"])
        aplicar_fonte_rodape(p_titulo)

        if meta.get("atualizacao"):
            p_atual = footer.add_paragraph(meta["atualizacao"])
            aplicar_fonte_rodape(p_atual)


def aplicar_layout_documento(doc, tipo_documento="GERAL"):
    # Página A4 com margens institucionais equilibradas.
    for section in doc.sections:
        section.top_margin = cm(3.2)
        section.bottom_margin = cm(2.0)
        section.left_margin = cm(2.3)
        section.right_margin = cm(2.0)
        section.header_distance = cm(1.0)
        section.footer_distance = cm(1.0)

    try:
        normal = doc.styles["Normal"]
        normal.font.name = FONTE_PADRAO
        normal._element.rPr.rFonts.set(qn('w:eastAsia'), FONTE_PADRAO)
        normal.font.size = Pt(TAM_BODY)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.08
    except Exception:
        pass

    # Parágrafos
    for p in doc.paragraphs:
        texto = p.text.strip()
        bold_hint = any(r.bold for r in p.runs)
        tipo = identificar_tipo_paragrafo(texto, bold_hint)
        # Preserva blocos de assinatura já centralizados no momento da geração.
        if p.alignment == WD_ALIGN_PARAGRAPH.CENTER and tipo not in {"titulo", "subtitulo", "secao"}:
            tipo = "assinatura"
        formatar_paragrafo(p, tipo)

    # Tabelas
    for t in doc.tables:
        try:
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.autofit = False
        except Exception:
            pass
        for row in t.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST if 'WD_ROW_HEIGHT_RULE' in globals() else row.height_rule
            for cell in row.cells:
                set_cell_text_format(cell)

    aplicar_rodape_modelo_atual(doc, tipo_documento)


def aplicar_larguras_tabela_objeto(table):
    larguras = [cm(1.0), cm(7.0), cm(2.3), cm(1.2), cm(2.4), cm(2.4)]
    for row_idx, row in enumerate(table.rows):
        for i, cell in enumerate(row.cells):
            if i < len(larguras):
                set_cell_width(cell, larguras[i])
            set_cell_text_format(cell, bold=(row_idx == 0), font_size=9 if i == 1 else 10)
            if row_idx == 0:
                set_cell_shading(cell, "D9EAF7")


def aplicar_larguras_tabela_area(table):
    larguras = [cm(5.0), cm(11.0)]
    for row_idx, row in enumerate(table.rows):
        for i, cell in enumerate(row.cells):
            if i < len(larguras):
                set_cell_width(cell, larguras[i])
            set_cell_text_format(cell, bold=(i == 0), font_size=10)
            if i == 0:
                set_cell_shading(cell, "EDEDED")


# =========================
# UTILITÁRIOS GERAIS
# =========================

def val(v):
    return str(v).strip() if v and str(v).strip() else COMPLEMENTAR


def clean(v):
    return "" if v is None else str(v).strip()


def val_assinatura(v, label=""):
    v = clean(v)
    if v:
        return v
    return (label + " ") if label else ""

def linha_campo(rotulo):
    return f"{rotulo}: __________________________________"


def normalizar_texto(s):
    if not s:
        return s
    return (
        str(s)
        .replace("Aniversario Da Cidade", "Aniversário da Cidade")
        .replace("Aniversario da Cidade", "Aniversário da Cidade")
        .replace("Aniversário Da Cidade", "Aniversário da Cidade")
        .replace("licitaçórios", "licitatórios")
        .replace("alvarsá", "alvarás")
    )




MESES_PT = {
    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril",
    5: "Maio", 6: "Junho", 7: "Julho", 8: "Agosto",
    9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro",
}


def normalizar_data_br(data):
    """
    Normaliza datas de entrada para dd/mm/aaaa quando possível.
    Aceita dd/mm/aaaa, dd-mm-aa, dd.mm.aa e datas por extenso.
    """
    if not data:
        return ""
    s = str(data).strip()
    if not s or s.lower() in ["a definir", "a informar", "não localizado", "nao localizado"]:
        return ""
    m = re.search(r"\b([0-3]?\d)[/.-]([01]?\d)[/.-]((?:20)?\d{2})\b", s)
    if m:
        d, mo, y = m.groups()
        if len(y) == 2:
            y = "20" + y
        return f"{int(d):02d}/{int(mo):02d}/{int(y):04d}"
    return s


def data_por_extenso(data):
    """
    Converte dd/mm/aaaa para '15 de Maio de 2026'.
    Se já estiver por extenso, mantém como informado.
    """
    s = normalizar_data_br(data)
    if not s:
        return ""
    m = re.match(r"^([0-3]\d)/([01]\d)/(20\d{2})$", s)
    if not m:
        return s
    d, mo, y = m.groups()
    return f"{int(d)} de {MESES_PT.get(int(mo), mo)} de {y}"


def data_atual_extenso():
    hoje = datetime.now()
    return f"{hoje.day} de {MESES_PT.get(hoje.month, hoje.month)} de {hoje.year}"


def data_documento_final(dados):
    # Fonte primária: data de emissão da Requisição. Fallback: data atual, nunca data fixa antiga.
    return data_por_extenso((dados or {}).get("data_documento")) or data_atual_extenso()


def data_evento_final(dados):
    return normalizar_data_br((dados or {}).get("data_evento")) or "a definir"


def data_conclusao_final(dados):
    # Se houver data de conclusão explicitamente informada, usa. Senão, usa a data do evento quando localizada.
    return normalizar_data_br((dados or {}).get("data_conclusao")) or normalizar_data_br((dados or {}).get("data_evento")) or "a definir"


def formatar_moeda(valor):
    if not valor:
        return ""
    s = str(valor).strip()
    if "R$" not in s:
        s = "R$ " + s
    s = s.replace("R$", "R$ ").replace("R$  ", "R$ ")
    return s


def moeda_para_numero(valor):
    if not valor:
        return 0.0
    s = str(valor)
    s = s.replace("R$", "").replace(" ", "").replace(".", "").replace(",", ".")
    try:
        return float(s)
    except Exception:
        return 0.0


UNIDADES = ["", "um", "dois", "três", "quatro", "cinco", "seis", "sete", "oito", "nove"]
DEZ_A_DEZENOVE = ["dez", "onze", "doze", "treze", "quatorze", "quinze", "dezesseis", "dezessete", "dezoito", "dezenove"]
DEZENAS = ["", "", "vinte", "trinta", "quarenta", "cinquenta", "sessenta", "setenta", "oitenta", "noventa"]
CENTENAS = ["", "cento", "duzentos", "trezentos", "quatrocentos", "quinhentos", "seiscentos", "setecentos", "oitocentos", "novecentos"]


def int_extenso(n):
    n = int(n)
    if n == 0:
        return "zero"
    if n == 100:
        return "cem"
    if n < 10:
        return UNIDADES[n]
    if n < 20:
        return DEZ_A_DEZENOVE[n - 10]
    if n < 100:
        dez, un = divmod(n, 10)
        return DEZENAS[dez] + ((" e " + UNIDADES[un]) if un else "")
    if n < 1000:
        cen, resto = divmod(n, 100)
        return CENTENAS[cen] + ((" e " + int_extenso(resto)) if resto else "")
    if n < 1000000:
        mil, resto = divmod(n, 1000)
        prefix = "mil" if mil == 1 else int_extenso(mil) + " mil"
        return prefix + ((" e " + int_extenso(resto)) if resto and resto < 100 else (" " + int_extenso(resto) if resto else ""))
    milhao, resto = divmod(n, 1000000)
    prefix = "um milhão" if milhao == 1 else int_extenso(milhao) + " milhões"
    return prefix + ((" e " + int_extenso(resto)) if resto else "")


def moeda_extenso(valor):
    n = moeda_para_numero(valor)
    inteiro = int(n)
    centavos = int(round((n - inteiro) * 100))
    if inteiro == 0:
        base = "zero real"
    elif inteiro == 1:
        base = "um real"
    else:
        base = int_extenso(inteiro) + " reais"
    if centavos:
        base += " e " + int_extenso(centavos) + (" centavo" if centavos == 1 else " centavos")
    return base


def data_padrao(dados):
    return dados.get("data_documento") or datetime.now().strftime("%d de %m de %Y")


def numero_processo(processo_id, dados):
    if dados.get("numero_processo"):
        return dados.get("numero_processo")
    if str(processo_id) not in ("", "None"):
        return f"Processo local nº {int(processo_id):04d}" if str(processo_id).isdigit() else str(processo_id)
    return "A informar"


def numero_dfd(processo_id, dados):
    # O número do DFD pode ser distinto do processo e da requisição.
    # Se não for informado pelo usuário ou extraído de documento próprio, manter campo a informar.
    return dados.get("numero_dfd") or "........."



def normalizar_nome_pessoa(nome):
    nome = clean(nome)
    nome = re.sub(r"\s+", " ", nome)
    nome = nome.strip(" -:;")
    # Corrige sobra herdada do modelo oficial quando havia "Matioli" no nome-base.
    if nome.upper().startswith("CRISTIAN MARCELO SCHIBELSKY"):
        return "Cristian Marcelo Schibelsky"
    return nome.title() if nome.isupper() else nome


def aplicar_cadastro_agente(d):
    nome_base = d.get("agente") or d.get("emissor_requisicao") or ""
    chave = clean(nome_base).upper()
    dados_agente = AGENTES_CONHECIDOS.get(chave)
    if not dados_agente and "CRISTIAN MARCELO SCHIBELSKY" in chave:
        dados_agente = AGENTES_CONHECIDOS["CRISTIAN MARCELO SCHIBELSKY"]

    if dados_agente:
        d["agente"] = dados_agente["nome"]
        if not d.get("agente_matricula"):
            d["agente_matricula"] = dados_agente["matricula"]
        if not d.get("agente_cargo"):
            d["agente_cargo"] = dados_agente["cargo"]
    else:
        d["agente"] = normalizar_nome_pessoa(d.get("agente"))
    return d


def enriquecer(dados, extraidos):
    d = {k: normalizar_texto(v) for k, v in dict(dados or {}).items()}
    extraidos = extraidos or {}
    d["descricao"] = normalizar_texto(d.get("descricao") or extraidos.get("objeto_sugerido", ""))
    d["valor_estimado"] = formatar_moeda(d.get("valor_estimado") or extraidos.get("valor_principal", ""))
    d["valor_extenso"] = d.get("valor_extenso") or (moeda_extenso(d["valor_estimado"]) if d["valor_estimado"] else "")
    d["origem_valor"] = d.get("origem_valor") or extraidos.get("origem_valor", "")
    d["artista"] = normalizar_texto(extraidos.get("artista", "")) or d.get("artista", "")
    d["cnpj"] = extraidos.get("cnpj", "")
    d["evento"] = normalizar_texto(extraidos.get("evento", "")) or d.get("evento", "")
    d["evidencias"] = extraidos.get("evidencias", {})
    d["valores_localizados"] = extraidos.get("valores_localizados", [])
    # Dados processuais/orçamentários extraídos automaticamente da Requisição, DFD, ETP, TR ou declaração orçamentária.
    for campo in [
        "numero_processo", "numero_requisicao", "numero_dfd", "codigo_reduzido", "funcional_programatica",
        "orgao", "unidade_orcamentaria", "acao", "subelemento", "vinculo", "condicao_pagamento", "local_entrega"
    ]:
        if not d.get(campo) and extraidos.get(campo):
            d[campo] = extraidos.get(campo)
    if not d.get("fonte_recurso") and extraidos.get("dotacao_texto"):
        d["fonte_recurso"] = extraidos.get("dotacao_texto")

    # Assinaturas e responsáveis extraídos da Requisição ao Compras.
    # Entram como sugestão automática e continuam editáveis no formulário.
    if not d.get("secretaria") and extraidos.get("secretaria_extraida"):
        d["secretaria"] = extraidos.get("secretaria_extraida")
    # Regra do projeto: o agente dos termos de designação ETP/TR deve ser o agente/emissor da Requisição.
    if extraidos.get("emissor_requisicao"):
        d["agente"] = extraidos.get("emissor_requisicao")
    elif not d.get("agente") and extraidos.get("emissor_requisicao"):
        d["agente"] = extraidos.get("emissor_requisicao")
    d = aplicar_cadastro_agente(d)
    if not d.get("responsavel_demanda") and extraidos.get("autoridade_assinante"):
        d["responsavel_demanda"] = extraidos.get("autoridade_assinante")
    if not d.get("secretario") and extraidos.get("autoridade_assinante"):
        d["secretario"] = extraidos.get("autoridade_assinante")
    if not d.get("cargo_secretario") and extraidos.get("cargo_autoridade"):
        d["cargo_secretario"] = extraidos.get("cargo_autoridade")
    if not d.get("responsavel_requisicao") and extraidos.get("responsavel_requisicao"):
        d["responsavel_requisicao"] = extraidos.get("responsavel_requisicao")

    # Regra definida: gestor e fiscal do contrato para estes processos.
    d["gestor"] = GESTOR_PADRAO
    d["gestor_cargo"] = GESTOR_CARGO_PADRAO
    d["gestor_competencias"] = GESTOR_COMPETENCIAS_PADRAO
    d["fiscal"] = FISCAL_PADRAO
    d["fiscal_cargo"] = FISCAL_CARGO_PADRAO
    d["fiscal_competencias"] = FISCAL_COMPETENCIAS_PADRAO
    return d


# =========================
# MANIPULAÇÃO DE DOCUMENTO
# =========================

def copiar_modelo(nome_modelo):
    caminho = MODELOS / nome_modelo
    return Document(str(caminho)) if caminho.exists() else Document()


def limpar_body_preservando_secao(doc):
    body = doc._element.body
    for child in list(body):
        if not child.tag.endswith("sectPr"):
            body.remove(child)



# =========================
# REVISÃO GRAMATICAL / ARTISTA
# =========================

ARTISTAS_MASCULINOS_CONHECIDOS = [
    "thiago brado",
    "luan pereira",
    "justin o frança",
    "justino frança",
    "padre",
    "cantor ",
]

ARTISTAS_FEMININOS_CONHECIDOS = [
    "cassiane",
    "gabriela rocha",
    "sarah farias",
    "sara farias",
    "aline barros",
    "bruna karla",
    "fernanda brum",
    "cantora ",
]


def genero_artista_por_nome(nome):
    n = (nome or "").strip().lower()
    if not n:
        return "neutro"
    if any(m in n for m in ARTISTAS_MASCULINOS_CONHECIDOS):
        return "masculino"
    if any(f in n for f in ARTISTAS_FEMININOS_CONHECIDOS):
        return "feminino"
    return "neutro"


def termos_artista(dados):
    artista_raw = clean((dados or {}).get("artista"))
    if not artista_raw or artista_raw.lower() in ["artista indicada", "artista indicado", "artista", "não localizado", "nao localizado"]:
        artista = ""
        genero = "neutro"
    else:
        artista = artista_raw
        genero = genero_artista_por_nome(artista)

    if genero == "masculino":
        return {
            "genero": genero,
            "profissao": "cantor",
            "profissao_artigo": "o cantor",
            "profissao_contracao": "do cantor",
            "artista_artigo": "o artista",
            "artista_contracao": "do artista",
            "qualificado": f"cantor {artista}",
            "contracao_qualificada": f"do cantor {artista}",
            "artigo_qualificado": f"o cantor {artista}",
        }
    if genero == "feminino":
        return {
            "genero": genero,
            "profissao": "cantora",
            "profissao_artigo": "a cantora",
            "profissao_contracao": "da cantora",
            "artista_artigo": "a artista",
            "artista_contracao": "da artista",
            "qualificado": f"cantora {artista}",
            "contracao_qualificada": f"da cantora {artista}",
            "artigo_qualificado": f"a cantora {artista}",
        }

    # Fallback limpo quando o artista ainda não foi identificado.
    return {
        "genero": genero,
        "profissao": "artista",
        "profissao_artigo": "o(a) artista",
        "profissao_contracao": "do(a) artista indicado(a)",
        "artista_artigo": "o(a) artista indicado(a)",
        "artista_contracao": "do(a) artista indicado(a)",
        "qualificado": "artista indicado(a)",
        "contracao_qualificada": "do(a) artista indicado(a)",
        "artigo_qualificado": "o(a) artista indicado(a)",
    }


def corrigir_gramatica_artistica_texto(texto, dados=None):
    """
    Revisor gramatical mínimo para evitar erros de concordância de gênero do artista,
    especialmente em contratações artísticas masculinas, como Thiago Brado.
    """
    if texto is None:
        return texto
    s = str(texto)

    nomes_masculinos = ["Thiago Brado", "Luan Pereira", "Justino França", "Justin o França"]
    for nome in nomes_masculinos:
        s = s.replace(f"da cantora {nome}", f"do cantor {nome}")
        s = s.replace(f"Da cantora {nome}", f"Do cantor {nome}")
        s = s.replace(f"a cantora {nome}", f"o cantor {nome}")
        s = s.replace(f"A cantora {nome}", f"O cantor {nome}")
        s = s.replace(f"cantora {nome}", f"cantor {nome}")
        s = s.replace(f"Cantora {nome}", f"Cantor {nome}")
        s = s.replace(f"da artista {nome}", f"do artista {nome}")
        s = s.replace(f"a artista {nome}", f"o artista {nome}")

    # Ajustes gerais de redação que apareceram nos modelos gerados.
    s = s.replace("da cantora artista indicada", "do(a) artista indicado(a)")
    s = s.replace("cantora artista indicada", "artista indicado(a)")
    s = s.replace("a artista artista indicada", "o(a) artista indicado(a)")
    s = s.replace("do(a) artista artista indicada", "do(a) artista indicado(a)")
    s = s.replace("do(a) artista artista indicado", "do(a) artista indicado(a)")
    s = s.replace("do(a) artista artista indicado(a)", "do(a) artista indicado(a)")
    s = s.replace("o(a) artista artista indicada", "o(a) artista indicado(a)")
    s = s.replace("o(a) artista artista indicado(a)", "o(a) artista indicado(a)")
    s = s.replace("artista artista indicada", "artista indicado(a)")
    s = s.replace("artista artista indicado(a)", "artista indicado(a)")
    s = s.replace("ao vivo da cantor", "ao vivo do cantor")
    s = s.replace("ao vivo da artista Thiago Brado", "ao vivo do artista Thiago Brado")
    s = s.replace("ao vivo do(a) artista artista indicada", "ao vivo do(a) artista indicado(a)")
    s = s.replace("apresentação artística ao vivo do(a) artista artista indicada", "apresentação artística ao vivo do(a) artista indicado(a)")

    # Concordância de preposição/artigo com profissão.
    s = s.replace("do cantora ", "da cantora ")
    s = s.replace("Do cantora ", "Da cantora ")
    s = s.replace("do cantorA ", "da cantora ")
    s = s.replace("ao vivo do cantora ", "ao vivo da cantora ")
    s = s.replace("apresentação artística ao vivo do cantora ", "apresentação artística ao vivo da cantora ")
    s = s.replace("realização de apresentação ao vivo do cantora ", "realização de apresentação ao vivo da cantora ")

    # Evita duplicidades de placeholder quando o artista ainda não foi identificado.
    s = s.replace("do(a) artista indicado(a) indicada", "do(a) artista indicado(a)")
    s = s.replace("artista indicado(a) indicada", "artista indicado(a)")
    s = s.replace("artista indicada indicada", "artista indicada")

    return s


def add_p(doc, texto="", style=None):
    p = doc.add_paragraph()
    tipo_formatacao = None
    if style in {"titulo", "subtitulo", "secao", "item", "assinatura", "campo", "body"}:
        tipo_formatacao = style
    elif style:
        try:
            p.style = style
        except Exception:
            pass
    texto = corrigir_gramatica_artistica_texto(texto)
    p.add_run(str(texto))
    formatar_paragrafo(p, tipo_formatacao or identificar_tipo_paragrafo(str(texto)))
    return p


def add_bold(doc, texto):
    p = doc.add_paragraph()
    texto = corrigir_gramatica_artistica_texto(texto)
    r = p.add_run(str(texto))
    r.bold = True
    formatar_paragrafo(p, identificar_tipo_paragrafo(str(texto), True))
    return p


def add_item(doc, texto):
    return add_p(doc, "· " + str(texto))


def add_assinatura(doc, cargo, nome=None, matricula=None):
    add_p(doc, "")
    add_p(doc, f"Sumaré, {data_atual_extenso()}.")
    add_p(doc, "")
    add_p(doc, "__________________________________")
    if nome:
        add_p(doc, nome)
    add_p(doc, cargo)
    if matricula:
        add_p(doc, f"Matrícula: {matricula}")


def add_bloco_assinatura(doc, data_doc, nome=None, cargo=None, matricula=None, rotulo=None):
    """Bloco de assinatura formal, centralizado e com espaçamento controlado.
    Preserva o modelo institucional e evita quebra/aperto entre assinatura e aprovação.
    """
    add_p(doc, "", style="assinatura")
    add_p(doc, f"Sumaré, {val(data_doc)}.", style="assinatura")
    add_p(doc, "", style="assinatura")
    add_p(doc, "_______________________________________", style="assinatura")
    if nome:
        add_p(doc, str(nome).upper(), style="assinatura")
    if matricula is not None:
        mat = clean(matricula)
        add_p(doc, f"Mat. {mat if mat else '__________________'}", style="assinatura")
    if cargo:
        add_p(doc, cargo, style="assinatura")
    if rotulo:
        add_p(doc, rotulo, style="assinatura")


def add_tabela_objeto(doc, descricao, valor):
    table = doc.add_table(rows=2, cols=6)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    headers = ["ITEM", "DESCRIÇÃO", "UNIDADE", "QTD.", "VALOR UNIT.", "VALOR TOTAL"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    row = table.rows[1].cells
    row[0].text = "1"
    row[1].text = corrigir_gramatica_artistica_texto(descricao)
    row[2].text = "Show / Apresentação"
    row[3].text = "1"
    row[4].text = valor
    row[5].text = valor
    aplicar_larguras_tabela_objeto(table)
    return table


def add_tabela_area_requisitante(doc, dados):
    table = doc.add_table(rows=5, cols=2)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    linhas = [
        ("Área Requisitante", val(dados.get("secretaria") or "SECRETARIA MUNICIPAL DE CULTURA E TURISMO")),
        ("Responsável pela demanda", val(dados.get("responsavel_demanda") or dados.get("secretario") or "")),
        ("matrícula", val(dados.get("matricula") or "")),
        ("Email", val(dados.get("email") or "smct@sumare.sp.gov.br")),
        ("Telefone", val(dados.get("telefone") or "(19) 3873-9469")),
    ]
    for i, (a, b) in enumerate(linhas):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
    aplicar_larguras_tabela_area(table)
    return table


def quadro_validacao_placeholders(doc):
    texto = "\n".join([p.text for p in doc.paragraphs])
    proibidos = ["XXX", "[ÓRGÃO", "[A complementar]", "empresas de cachê", "notória especialização"]
    encontrados = [x for x in proibidos if x in texto]
    return encontrados


# =========================
# TEXTOS TÉCNICO-JURÍDICOS
# =========================

def objeto_padrao(dados):
    artista = val(dados.get("artista") or "artista indicada")
    evento = val(dados.get("evento") or "evento municipal")
    termos = termos_artista(dados)
    descricao = dados.get("descricao") or f"Contratação de serviços artísticos para a realização de apresentação ao vivo {termos['contracao_qualificada']}, no evento “{evento}”, a ser realizado no Município de Sumaré/SP"
    return corrigir_gramatica_artistica_texto(descricao, dados)


def texto_justificativa_dfd(dados):
    evento = val(dados.get("evento") or "evento municipal")
    termos = termos_artista(dados)
    return corrigir_gramatica_artistica_texto(
        f"A presente contratação tem por finalidade atender à necessidade da Secretaria Municipal de Cultura e Turismo de compor a programação do evento “{evento}”, "
        f"mediante apresentação artística ao vivo {termos['contracao_qualificada']}. A demanda possui pertinência com as atribuições institucionais da unidade requisitante, voltadas à promoção de atividades culturais, "
        "fomento ao turismo, valorização de eventos públicos municipais, ampliação do acesso da população a atividades de lazer e fortalecimento da convivência comunitária. "
        "A contratação deverá ser instruída com demonstração da razão da escolha do(a) artista, comprovação de representação exclusiva ou contratação direta, elementos de consagração pela crítica especializada "
        "ou pela opinião pública, justificativa de preço, disponibilidade orçamentária e demais documentos de habilitação exigíveis.",
        dados
    )


def texto_exclusividade(evidencias):
    ev = (evidencias or {}).get("exclusividade", {})
    arq = ev.get("arquivo")
    if ev.get("ok"):
        base = f"Conforme documentação de exclusividade juntada aos autos{(' (' + arq + ')') if arq else ''},"
        return (
            f"{base} a contratação deve ser analisada sob a perspectiva da inviabilidade de competição própria das apresentações artísticas, "
            "pois a Administração não busca serviço artístico genérico, mas apresentação de artista determinada. A validade dessa premissa exige conferência da autenticidade, vigência, abrangência territorial e temporal, "
            "identificação da artista, identificação da empresa representante e poderes de quem subscreve o documento."
        )
    return (
        "A contratação direta por inexigibilidade depende da juntada e validação de documento idôneo que comprove a contratação direta da artista ou a atuação de empresário exclusivo, "
        "com identificação da artista, representante, vigência, abrangência e poderes de representação. Sem esse elemento, a conclusão pela inviabilidade de competição fica prejudicada."
    )


def texto_consagracao(evidencias, artista):
    ev = (evidencias or {}).get("consagracao", {})
    arq = ev.get("arquivo")
    if ev.get("ok"):
        return (
            f"Os documentos de consagração pública ou crítica especializada juntados aos autos{(' (' + arq + ')') if arq else ''} devem ser considerados para demonstrar que a artista {artista} possui reconhecimento compatível com o porte e a finalidade do evento. "
            "A análise deve recair sobre elementos objetivos, tais como histórico artístico, registros de mídia, agenda de apresentações, alcance público, material de divulgação, repercussão no segmento musical e aderência entre o perfil artístico e o público-alvo do evento."
        )
    return (
        f"A instrução deverá ser complementada com elementos objetivos que demonstrem a consagração da artista {artista} pela crítica especializada ou pela opinião pública, tais como release, clipping, registros de mídia, histórico de apresentações, agenda, dados de audiência, materiais de divulgação e demais documentos equivalentes."
    )


def texto_preco(evidencias, dados):
    valor = val(dados.get("valor_estimado"))
    origem = val(dados.get("origem_valor") or "documentos de instrução")
    ev = (evidencias or {}).get("justificativa_preco", {})
    arq = ev.get("arquivo")
    apoio = f", especialmente {arq}," if ev.get("ok") and arq else ""
    return (
        f"O valor estimado da contratação é de {valor} ({val(dados.get('valor_extenso'))}), tendo como referência {origem}. "
        f"A justificativa de preço deve ser aferida a partir dos documentos constantes dos autos{apoio} observando-se que notas fiscais, contratos anteriores e contratações similares não substituem a proposta atual nem a requisição do processo, "
        "mas servem como parâmetros de controle de razoabilidade. A análise deve considerar data da contratação similar, ente contratante, localidade, objeto executado, porte do evento, condições de deslocamento, escopo assumido pela contratada e eventuais diferenças materiais entre as contratações comparadas."
    )


def texto_pca():
    return (
        "A unidade requisitante deverá certificar a previsão da contratação no Plano de Contratações Anual, calendário oficial de eventos, planejamento setorial ou instrumento equivalente. "
        "Na hipótese de ausência de previsão específica, deverá justificar a inclusão da demanda, demonstrando interesse público, oportunidade administrativa, compatibilidade orçamentária e aderência às finalidades da Secretaria."
    )


def texto_resultados(dados):
    evento = val(dados.get("evento") or "evento municipal")
    return (
        f"Com a realização do evento “{evento}”, a Administração pretende promover evento cultural de interesse público, ampliar o acesso da população a atividade artística, "
        "fortalecer o calendário municipal, fomentar a economia local de forma indireta e assegurar melhor aproveitamento dos recursos públicos empregados na programação, mediante contratação previamente planejada, motivada e fiscalizada."
    )


def dotacao_texto(dados):
    dot = clean(dados.get("fonte_recurso"))
    if dot:
        return dot
    return "A dotação orçamentária deverá ser indicada pela unidade competente antes da formalização final, preferencialmente com órgão, unidade, funcional programática, ação, elemento, subelemento, código reduzido, vínculo e fonte de recurso."


# =========================
# GERAÇÃO DFD
# =========================

def gerar_requisicao(processo_id, dados, classificacao, pasta_saida: Path, extraidos=None):
    # A requisição oficial é emitida pelo sistema Atende.Net/IPM. Gera-se apenas referência baseada no DFD para compatibilidade do app.
    return gerar_dfd(processo_id, dados, classificacao, pasta_saida, extraidos, nome_arquivo="01_DFD_base_requisicao.docx")


def gerar_dfd(processo_id, dados, classificacao, pasta_saida: Path, extraidos=None, nome_arquivo="02_DFD.docx"):
    dados = enriquecer(dados, extraidos)
    pasta_saida.mkdir(parents=True, exist_ok=True)
    arquivo = pasta_saida / nome_arquivo
    doc = copiar_modelo("DFD_Modelo_2025.docx")
    limpar_body_preservando_secao(doc)

    secretaria = dados.get("secretaria") or "SECRETARIA MUNICIPAL DE CULTURA E TURISMO"
    objeto = objeto_padrao(dados)
    valor = val(dados.get("valor_estimado"))
    valor_ext = val(dados.get("valor_extenso"))

    add_bold(doc, "DOCUMENTO DE FORMALIZAÇÃO DE DEMANDA")
    add_p(doc, f" nº {numero_dfd(processo_id, dados)}")
    add_p(doc, "")
    add_p(doc, secretaria.upper())

    add_bold(doc, "1. IDENTIFICAÇÃO DA DEMANDA")
    add_p(doc, "1.1. TIPO: SERVIÇO")
    add_p(doc, "1.2. CLASSE: EVENTOS")
    add_p(doc, f"1.3. OBJETO: {objeto}")
    add_p(doc, "Serviço ou fornecimento continuado?     (         )   SIM                  (     X     ) NÃO")

    add_bold(doc, "2. DESCRIÇÃO SUCINTA DO OBJETO")
    add_p(doc, objeto)

    add_bold(doc, "3. QUANTIDADE A SER CONTRATADA")
    add_p(doc, "01 apresentação/show")

    add_bold(doc, "4. ESTIMATIVA PRELIMINAR DO VALOR DA CONTRATAÇÃO")
    add_p(doc, f"{valor} ({valor_ext})")

    add_bold(doc, "5. JUSTIFICATIVA DA NECESSIDADE DA CONTRATAÇÃO")
    add_p(doc, texto_justificativa_dfd(dados))

    add_bold(doc, "6. INDICAÇÃO DE VINCULAÇÃO OU DEPENDÊNCIA COM O OBJETO DE OUTRO DOCUMENTO DE FORMALIZAÇÃO DE DEMANDA")
    add_p(doc, "O objeto desta contratação não possui dependência direta de outro Documento de Formalização de Demanda para a contratação artística principal. Todavia, a execução do evento poderá demandar contratações correlatas ou providências administrativas autônomas, tais como palco, sonorização, iluminação, segurança, banheiros químicos, limpeza, brigada, ambulância, camarim, logística, comunicação e demais estruturas de apoio, a serem tratadas nos instrumentos próprios, quando cabível.")

    add_bold(doc, "7. DATA PRETENDIDA PARA A CONCLUSÃO DA CONTRATAÇÃO:")
    add_p(doc, f"até {val(data_conclusao_final(dados))}")
    add_p(doc, "Grau de prioridade da contratação:")
    add_p(doc, "(             ) baixa")
    add_p(doc, "(             ) média")
    add_p(doc, "(       X      ) alta")
    add_p(doc, "Justificativa da prioridade atribuída: contratação vinculada à programação de evento público, cuja execução exige instrução processual tempestiva, análise documental, autorização da autoridade competente e formalização prévia.")

    add_bold(doc, "8. ÁREA REQUISITANTE DA DEMANDA")
    add_tabela_area_requisitante(doc, dados)

    add_p(doc, "")
    add_p(doc, f"Sumaré, {val(data_documento_final(dados))}.", style="assinatura")
    add_p(doc, "__________________________________", style="assinatura")
    resp_nome = clean(dados.get("responsavel_demanda") or dados.get("secretario"))
    cargo_resp = clean(dados.get("cargo_secretario") or secretaria)
    add_p(doc, resp_nome.upper() if resp_nome else "__________________________________", style="assinatura")
    add_p(doc, cargo_resp.upper() if cargo_resp else "CARGO/FUNÇÃO A INFORMAR", style="assinatura")

    aplicar_layout_documento(doc, "DFD")
    doc.save(arquivo)
    return arquivo


# =========================
# GERAÇÃO ETP
# =========================

def gerar_etp(processo_id, dados, classificacao, pasta_saida: Path, extraidos=None):
    dados = enriquecer(dados, extraidos)
    pasta_saida.mkdir(parents=True, exist_ok=True)
    arquivo = pasta_saida / "03_ETP.docx"
    doc = copiar_modelo("ETP_Modelo_2025.docx")
    limpar_body_preservando_secao(doc)

    secretaria = dados.get("secretaria") or "SECRETARIA MUNICIPAL DE CULTURA E TURISMO"
    objeto = objeto_padrao(dados)
    artista = val(dados.get("artista") or "artista indicada")
    evento = val(dados.get("evento") or "evento municipal")
    evid = dados.get("evidencias", {})

    add_bold(doc, "ESTUDO TÉCNICO PRELIMINAR")
    add_p(doc, secretaria.upper())
    add_p(doc, f"(Processo Administrativo n° {numero_processo(processo_id, dados)})")

    add_bold(doc, "1. DESCRIÇÃO DA NECESSIDADE DA CONTRATAÇÃO – art. 18, I c/c §1º, I")
    add_p(doc, f"O Município de Sumaré, por meio da {secretaria}, pretende realizar o evento denominado “{evento}”, voltado à população do Município, identificando a necessidade de contratação de apresentação artística ao vivo para compor sua programação principal.")
    add_p(doc, f"A contratação {termos_artista(dados)['contracao_qualificada']} deve ser analisada sob a perspectiva do interesse público cultural, da adequação da atração ao perfil do evento, da motivação administrativa, da demonstração de consagração pela crítica especializada ou pela opinião pública e da comprovação da inviabilidade de competição nos termos do art. 74, inciso II, da Lei Federal nº 14.133/2021.")
    add_p(doc, "A necessidade pública não se resume à mera realização de show, mas à viabilização de programação cultural planejada, com acesso da população a atividade artística de relevância, observada a adequada instrução processual, a disponibilidade orçamentária e a fiscalização da execução.")

    add_bold(doc, "2. PREVISÃO DA CONTRATAÇÃO NO PLANO DE CONTRATAÇÕES ANUAL, SE HOUVER - art. 18, I")
    add_p(doc, texto_pca())

    add_bold(doc, "3. REQUISITOS DA CONTRATAÇÃO - art. 18, III")
    add_p(doc, "3.1. Os requisitos da contratação serão pormenorizados no Termo de Referência, sem prejuízo dos requisitos mínimos identificados neste Estudo Técnico Preliminar:")
    for item in [
        "demonstração da necessidade pública e da pertinência da atração com o evento;",
        "comprovação de contratação direta da artista ou de representação por empresário exclusivo;",
        "comprovação de consagração pela crítica especializada ou pela opinião pública;",
        "justificativa de preço com parâmetros de compatibilidade e documentos de suporte;",
        "regularidade jurídica, fiscal, social e trabalhista da futura contratada, conforme aplicável;",
        "definição de data, local, horário, condições de execução, gestor e fiscal;",
        "verificação das contratações correlatas necessárias à execução segura e adequada do evento."
    ]:
        add_item(doc, item)

    add_bold(doc, "4. ESTIMATIVA DA QUANTIDADE – art. 18, IV")
    add_p(doc, f"Estima-se a contratação de 01 (uma) apresentação artística/show, por se tratar de execução única vinculada ao evento “{evento}”. A quantidade decorre da própria natureza da prestação pretendida, não havendo demanda por quantitativos recorrentes ou fornecimento continuado.")

    add_bold(doc, "5. LEVANTAMENTO DE MERCADO – art. 18, V")
    add_p(doc, "O levantamento de mercado, em contratações artísticas por inexigibilidade, não tem por finalidade comparar artistas como se fossem bens ou serviços comuns substituíveis. A análise deve verificar se a contratação de artista determinada encontra respaldo na consagração pública ou crítica especializada, na exclusividade de representação e na compatibilidade do preço com contratações anteriores ou similares.")
    add_p(doc, texto_exclusividade(evid))
    add_p(doc, texto_consagracao(evid, artista))

    add_bold(doc, "6. ESTIMATIVA DO VALOR DA CONTRATAÇÃO – art. 18, VI")
    add_p(doc, texto_preco(evid, dados))

    add_bold(doc, "7. DESCRIÇÃO DA SOLUÇÃO – art. 18, VII")
    add_p(doc, f"A solução consiste na realização de apresentação artística ao vivo {termos_artista(dados)['contracao_qualificada']} durante o evento “{evento}”, no Município de Sumaré/SP, conforme condições a serem definidas no Termo de Referência, na proposta e nos demais documentos de instrução.")
    for item in [
        "apresentação artística ao vivo, com duração mínima a ser definida no Termo de Referência ou proposta;",
        "equipe artística e técnica própria da contratada, quando aplicável;",
        "observância ao rider técnico, camarim, logística e demais condições pactuadas;",
        "estrutura de palco, sonorização, iluminação, segurança e demais itens de apoio a cargo do Município ou de contratações correlatas, salvo disposição expressa em sentido diverso."
    ]:
        add_item(doc, item)

    add_bold(doc, "8. JUSTIFICATIVAS PARA O PARCELAMENTO OU NÃO DA CONTRATAÇÃO – art. 18, VIII")
    add_p(doc, "O objeto da contratação artística principal é indivisível por natureza, pois consiste em apresentação única de artista determinada. O parcelamento da apresentação descaracterizaria a finalidade pública e a identidade artística pretendida. Serviços acessórios e autônomos necessários à execução do evento poderão ser objeto de contratações próprias, quando técnica e administrativamente cabível.")

    add_bold(doc, "9. RESULTADOS PRETENDIDOS – art. 18, IX")
    add_p(doc, texto_resultados(dados))

    add_bold(doc, "10. PROVIDÊNCIAS A SEREM ADOTADAS PELA ADMINISTRAÇÃO PREVIAMENTE AO CONTRATO– art. 18, X")
    for item in [
        "confirmar data, horário e local do evento;",
        "verificar disponibilidade orçamentária e fonte de recurso;",
        "validar documento de exclusividade quanto à autenticidade, vigência, abrangência e poderes de representação;",
        "conferir documentos de consagração pública ou crítica especializada;",
        "consolidar justificativa de preço com documentos comparativos;",
        "designar gestor e fiscal do contrato;",
        "verificar infraestrutura, segurança, acessibilidade, licenças, alvarás, camarim, logística e demais condições operacionais;",
        "submeter a instrução à análise jurídica e à autorização da autoridade competente."
    ]:
        add_item(doc, item)

    add_bold(doc, "11. CONTRATAÇÕES CORRELATAS / INTERDEPENDENTES – art. 18, XI")
    add_p(doc, "A execução do evento poderá exigir contratações correlatas ou providências administrativas autônomas, tais como locação e montagem de palco, sonorização, iluminação, gerador, segurança, controle de acesso, limpeza, sanitários móveis, ambulância, brigada, camarim, comunicação institucional, autorizações e demais serviços de apoio. Tais providências deverão observar os procedimentos próprios e os prazos compatíveis com a realização do evento.")

    add_bold(doc, "12. IMPACTOS AMBIENTAIS E TRATAMENTOS")
    add_p(doc, "A contratação artística principal não apresenta, isoladamente, impacto ambiental significativo. A realização do evento, contudo, poderá gerar resíduos, ruído e fluxo intensificado de pessoas, recomendando-se medidas de limpeza, organização do espaço, controle de resíduos, comunicação preferencialmente digital quando possível e observância das normas de segurança, acessibilidade e proteção do patrimônio público.")

    add_bold(doc, "13. VIABILIDADE DA CONTRATAÇÃO")
    add_p(doc, "Diante dos elementos analisados, conclui-se pela viabilidade condicionada da contratação direta por inexigibilidade de licitação, com fundamento no art. 74, inciso II, da Lei Federal nº 14.133/2021, desde que, antes da formalização final, estejam comprovados nos autos:")
    for item in [
        "a razão da escolha da artista e sua pertinência com o evento;",
        "a contratação direta ou representação por empresário exclusivo;",
        "a consagração pela crítica especializada ou pela opinião pública;",
        "a justificativa de preço com parâmetros idôneos;",
        "a disponibilidade orçamentária;",
        "a regularidade documental da futura contratada;",
        "a definição das condições de execução, gestão e fiscalização."
    ]:
        add_item(doc, item)

    add_bold(doc, "14. CLASSIFICAÇÃO DA INFORMAÇÃO")
    add_p(doc, "14.1. Conforme disposto no Decreto Municipal aplicável, não se identifica, nesta minuta, conteúdo que demande classificação de sigilo, sem prejuízo da avaliação pela unidade competente quanto à existência de dados pessoais, informações estratégicas ou anexos que eventualmente exijam tratamento específico.")

    agente_nome = clean(dados.get("agente"))
    agente_matricula = clean(dados.get("agente_matricula"))
    add_bloco_assinatura(
        doc,
        data_documento_final(dados),
        nome=agente_nome if agente_nome else None,
        matricula=agente_matricula,
        rotulo="Agente responsável pela elaboração do Estudo Técnico Preliminar - ETP",
    )

    doc.add_page_break()
    add_bold(doc, "TERMO DE APROVAÇÃO DE ESTUDO TÉCNICO PRELIMINAR")
    add_p(doc, "Nos termos do Decreto Municipal nº 12.053/2023, APROVO e FIRMO o respectivo Estudo Técnico Preliminar - ETP, uma vez que a solução apresentada, observadas as condições e complementações indicadas, mostra-se apta a satisfazer o interesse público.")
    autoridade_nome = clean(dados.get("secretario"))
    cargo_aut = clean(dados.get("cargo_secretario"))
    add_bloco_assinatura(
        doc,
        data_documento_final(dados),
        nome=autoridade_nome if autoridade_nome else None,
        cargo=cargo_aut if cargo_aut else "Cargo da autoridade aprovadora: ______________________________",
    )

    add_bold(doc, "REFERÊNCIAS:")
    add_p(doc, "Base legal: Lei Federal nº 14.133/2021, especialmente arts. 6º, XX; 18; 72; 74, inciso II; e demais dispositivos aplicáveis à fase preparatória e à contratação direta.")
    add_p(doc, "Decretos Municipais aplicáveis à fase preparatória, designação de agentes, elaboração e aprovação do Estudo Técnico Preliminar e classificação da informação.")

    aplicar_layout_documento(doc, "ETP")
    doc.save(arquivo)
    return arquivo


# =========================
# GERAÇÃO TR
# =========================

def gerar_tr(processo_id, dados, classificacao, pasta_saida: Path, extraidos=None):
    dados = enriquecer(dados, extraidos)
    pasta_saida.mkdir(parents=True, exist_ok=True)
    arquivo = pasta_saida / "04_TR.docx"
    doc = copiar_modelo("TR_Servicos_Sem_Demo_2025.docx")
    limpar_body_preservando_secao(doc)

    secretaria = dados.get("secretaria") or "SECRETARIA MUNICIPAL DE CULTURA E TURISMO"
    objeto = objeto_padrao(dados)
    artista = val(dados.get("artista") or "artista indicada")
    evento = val(dados.get("evento") or "evento municipal")
    valor = val(dados.get("valor_estimado"))
    valor_ext = val(dados.get("valor_extenso"))
    evid = dados.get("evidencias", {})

    add_bold(doc, "TERMO DE REFERÊNCIA")
    add_p(doc, "Serviços sem dedicação exclusiva de mão-de-obra")
    add_p(doc, "Licitação ou Contratação Direta")
    add_p(doc, secretaria.upper())
    add_p(doc, f"(Processo Administrativo n° {numero_processo(processo_id, dados)})")

    add_bold(doc, "1. FUNDAMENTAÇÃO DA NECESSIDADE DA CONTRATAÇÃO – Art. 6º, XXIII, “b”")
    add_p(doc, "1.1. A fundamentação da contratação encontra-se pormenorizada no Estudo Técnico Preliminar, apêndice deste Termo de Referência.")
    add_p(doc, f"1.2. A contratação visa atender à necessidade da {secretaria} de compor a programação do evento “{evento}” com apresentação artística ao vivo {termos_artista(dados)['contracao_qualificada']}, observada a finalidade pública de promoção cultural, lazer, integração comunitária e valorização do calendário municipal.")
    add_p(doc, "1.3. O enquadramento preliminar é de contratação direta por inexigibilidade de licitação, com fundamento no art. 74, inciso II, da Lei Federal nº 14.133/2021, condicionado à comprovação de que se trata de profissional do setor artístico, contratado diretamente ou por empresário exclusivo, consagrado pela crítica especializada ou pela opinião pública.")

    add_bold(doc, "2. DEFINIÇÃO DO OBJETO – Art. 6º, XXIII, “a” e Art. 40, §1º, I e III")
    add_p(doc, f"2.1. {objeto}, conforme tabela abaixo:")
    add_tabela_objeto(doc, f"Apresentação artística ao vivo {termos_artista(dados)['contracao_qualificada']}", valor)
    add_p(doc, "2.2. O serviço objeto desta contratação é caracterizado como serviço artístico, sem dedicação exclusiva de mão de obra, contratado por escopo, com padrões de execução definidos pelo Termo de Referência, proposta e demais anexos.")
    add_p(doc, "2.3. O serviço é enquadrado como não contínuo, pois se exaure com a realização da apresentação e das providências necessárias ao recebimento, liquidação e pagamento.")
    add_p(doc, f"2.4. O prazo de vigência da contratação deverá ser suficiente para formalização, execução, recebimento, liquidação e pagamento do objeto, observados os créditos orçamentários aplicáveis. Data/local da execução: {val(data_evento_final(dados))}; {val(dados.get('local_evento') or 'local a definir pela unidade requisitante')}.")

    add_bold(doc, "3. DESCRIÇÃO DA SOLUÇÃO - Art. 6º, XXIII, “c”")
    add_p(doc, "3.1. A descrição da solução como um todo encontra-se pormenorizada em tópico específico do Estudo Técnico Preliminar.")
    add_p(doc, f"3.2. A solução consiste na contratação da apresentação artística ao vivo {termos_artista(dados)['contracao_qualificada']}, no evento “{evento}”, compreendendo a prestação artística principal, equipe artística e técnica própria quando aplicável, observância às condições pactuadas e cumprimento da proposta aprovada.")
    add_p(doc, "3.3. Não integram automaticamente o presente objeto os serviços de palco, som, iluminação, segurança, banheiros, brigada, ambulância, limpeza, ECAD, camarim, hospedagem, alimentação ou logística, salvo quando expressamente previstos na proposta ou no instrumento contratual.")

    add_bold(doc, "4. ESTIMATIVA DO VALOR DA CONTRATAÇÃO - Art. 6º, XXIII, “i”")
    add_p(doc, f"4.1. O custo estimado total da contratação é de {valor} ({valor_ext}), conforme valor de referência extraído de {val(dados.get('origem_valor') or 'documentos de instrução')} e documentos de suporte constantes dos autos.")
    add_p(doc, "4.2. A estimativa de custo deverá ser acompanhada de justificativa de preço, memória de análise ou documentos comparativos que permitam aferir a razoabilidade do valor.")
    add_p(doc, "4.3. Para fins de controle, notas fiscais, contratos anteriores e contratações similares constituem parâmetros de compatibilidade, devendo ser avaliados quanto à data, contratante, localidade, porte do evento, escopo, artista e condições de execução.")

    add_bold(doc, "5. ADEQUAÇÃO ORÇAMENTÁRIA - Art. 6º, XXIII, “j”")
    add_p(doc, "5.1. As despesas decorrentes da presente contratação correrão à conta de recursos específicos consignados no Orçamento Geral do Município de Sumaré.")
    add_p(doc, f"5.2. A contratação será atendida pela seguinte dotação/fonte de recurso: {dotacao_texto(dados)}")

    add_bold(doc, "6. FORMA E CRITÉRIOS DE SELEÇÃO DO FORNECEDOR - art. 6º, XXIII, “h”")
    add_p(doc, "6.1. O fornecedor será selecionado por meio de contratação direta, por inexigibilidade de licitação, com fundamento no art. 74, inciso II, da Lei Federal nº 14.133/2021, em razão da contratação de profissional do setor artístico, diretamente ou por meio de empresário exclusivo, desde que comprovada a consagração pela crítica especializada ou pela opinião pública.")
    add_p(doc, "6.2. O regime de execução do serviço será de empreitada por preço global, considerando a execução integral da apresentação artística contratada.")

    add_bold(doc, "Exigências de habilitação")
    habilitacao = [
        "prova de inscrição no Cadastro Nacional da Pessoa Jurídica ou Cadastro de Pessoa Física, conforme o caso;",
        "ato constitutivo, contrato social, estatuto, alterações ou documento equivalente, acompanhado de comprovação de poderes de representação;",
        "regularidade fiscal perante a Fazenda Nacional;",
        "regularidade com o Fundo de Garantia do Tempo de Serviço, quando aplicável;",
        "regularidade perante a Justiça do Trabalho;",
        "regularidade municipal pertinente ao ramo de atividade, quando aplicável;",
        "documento de exclusividade ou representação artística idôneo;",
        "documentos de consagração pública ou crítica especializada;",
        "documentos de suporte à justificativa de preço;",
        "declarações exigidas pela legislação e pelos modelos municipais aplicáveis."
    ]
    for idx, item in enumerate(habilitacao, start=1):
        add_p(doc, f"6.3.{idx}. {item}")

    add_bold(doc, "7. REQUISITOS DA CONTRATAÇÃO - Art. 6º, XXIII, “d”")
    add_p(doc, "7.1. Sustentabilidade: os critérios ambientais e de sustentabilidade devem observar o quanto disposto no Estudo Técnico Preliminar e no planejamento operacional do evento.")
    add_p(doc, "7.2. Subcontratação: não será admitida a subcontratação da apresentação artística principal, por incompatibilidade com a natureza personalíssima do objeto.")
    add_p(doc, "7.3. Garantia da contratação: não haverá exigência de garantia contratual, salvo decisão administrativa fundamentada em sentido diverso.")
    add_p(doc, "7.4. Vistoria: não há necessidade de vistoria prévia pela contratada para a apresentação artística, sem prejuízo de alinhamento técnico quanto a palco, sonorização, iluminação, camarim e demais condições operacionais.")

    add_bold(doc, "8. MODELO DE EXECUÇÃO DO OBJETO - Art. 6º, XXIII, “e” e Art. 40, §1º, II")
    add_p(doc, f"8.1. A execução do objeto ocorrerá mediante apresentação artística ao vivo {termos_artista(dados)['contracao_qualificada']}, em data, horário e local a serem definidos ou confirmados pela unidade requisitante e pela contratada.")
    add_p(doc, f"8.2. Local e horário da prestação dos serviços: {val(dados.get('local_evento') or 'a definir pela unidade requisitante')}; {val(dados.get('horario') or 'horário a definir')}.")
    add_p(doc, "8.3. A duração mínima da apresentação, passagem de som, rider técnico, camarim, equipe artística, deslocamento, hospedagem, alimentação e demais condições específicas deverão constar da proposta, contrato ou anexo técnico validado pela Administração.")

    add_bold(doc, "9. DA GESTÃO DA CONTRATAÇÃO – Art. 6º, XXIII, “f”")
    add_p(doc, "9.1. O contrato ou instrumento equivalente deverá ser executado fielmente pelas partes, de acordo com as cláusulas avençadas e as normas da Lei Federal nº 14.133/2021.")
    add_p(doc, f"9.2. A gestão da contratação caberá a {val(dados.get('gestor') or 'gestor a ser designado formalmente')}, e a fiscalização caberá a {val(dados.get('fiscal') or 'fiscal a ser designado formalmente')}, ou aos respectivos substitutos formalmente indicados.")
    add_p(doc, "9.3. Compete à fiscalização acompanhar a execução, verificar a realização da apresentação, registrar ocorrências, conferir o cumprimento das condições pactuadas, atestar a execução e encaminhar documentos para liquidação e pagamento.")
    add_p(doc, "9.4. A fiscalização não exclui nem reduz a responsabilidade da contratada por falhas, descumprimentos, atrasos, inconsistências documentais ou inexecução total ou parcial.")

    add_bold(doc, "10. CRITÉRIOS DE RECEBIMENTO E PAGAMENTO - Art. 6º, XXIII, “g”")
    add_p(doc, "10.1. Os serviços serão recebidos provisoriamente pelo fiscal do contrato mediante registro da execução da apresentação artística e verificação do cumprimento das exigências de caráter técnico e administrativo.")
    add_p(doc, "10.2. O recebimento definitivo ficará condicionado à manifestação da autoridade competente ou responsável designado, com base nas informações do fiscal e do gestor do contrato.")
    add_p(doc, "10.3. Os serviços poderão ser rejeitados, no todo ou em parte, quando executados em desacordo com as especificações constantes deste Termo de Referência, proposta ou instrumento contratual.")
    prazo_pagamento = dados.get("condicao_pagamento") or "até 28 (vinte e oito) dias úteis a contar do recebimento da nota fiscal ou instrumento de cobrança equivalente pela Administração"
    add_p(doc, f"10.4. O pagamento será efetuado {prazo_pagamento}, devidamente atestada pela fiscalização e pelo gestor do contrato, observadas as normas municipais aplicáveis.")
    add_p(doc, "10.5. Para fins de liquidação, o setor competente deverá verificar a nota fiscal, o atesto, os dados do contrato, o valor a pagar, eventual retenção tributária cabível e a manutenção das condições de habilitação exigidas.")
    add_p(doc, "10.6. Havendo erro na nota fiscal, pendência documental ou circunstância que impeça a liquidação, o pagamento ficará suspenso até saneamento pela contratada, sem ônus ao contratante.")
    add_p(doc, "10.7. O pagamento será realizado por ordem bancária em conta indicada pela contratada, com retenções tributárias previstas na legislação aplicável.")

    add_bold(doc, "11. INFRAÇÕES, MULTAS E SANÇÕES")
    add_p(doc, "11.1. As infrações, multas e sanções aplicáveis à futura contratada serão aquelas previstas na Lei Federal nº 14.133/2021, no Decreto Municipal aplicável e no instrumento contratual, assegurados contraditório e ampla defesa.")
    add_p(doc, "11.2. A inexecução total ou parcial, ausência injustificada, descumprimento de horário, execução em desconformidade, irregularidade documental ou inadimplemento de obrigação assumida poderão ensejar apuração de responsabilidade e aplicação das medidas cabíveis.")

    add_bold(doc, "DISPOSIÇÕES FINAIS")
    add_p(doc, "Este Termo de Referência fará parte integrante do futuro contrato ou do instrumento hábil que o substitua nas hipóteses legais.")
    add_p(doc, "A autorização final da contratação fica condicionada à consolidação da instrução processual, comprovação da exclusividade, consagração, justificativa de preço, disponibilidade orçamentária, regularidade documental, designação de gestor e fiscal e análise jurídica, quando exigível.")

    agente_nome = clean(dados.get("agente"))
    agente_matricula = clean(dados.get("agente_matricula"))
    add_bloco_assinatura(
        doc,
        data_documento_final(dados),
        nome=agente_nome if agente_nome else None,
        matricula=agente_matricula,
        rotulo="Responsável pela elaboração do Termo de Referência",
    )

    doc.add_page_break()
    add_bold(doc, "TERMO DE APROVAÇÃO DE TERMO DE REFERÊNCIA")
    add_p(doc, "Nos termos do Decreto Municipal nº 12.054/2023, APROVO e FIRMO o respectivo Termo de Referência.")
    autoridade_nome = clean(dados.get("secretario"))
    cargo_aut = clean(dados.get("cargo_secretario"))
    add_bloco_assinatura(
        doc,
        data_documento_final(dados),
        nome=autoridade_nome if autoridade_nome else None,
        cargo=cargo_aut if cargo_aut else "Cargo da autoridade aprovadora: ______________________________",
    )

    aplicar_layout_documento(doc, "TR")
    doc.save(arquivo)
    return arquivo


# =========================
# DOCUMENTOS AUXILIARES
# =========================

def replace_doc_text(doc, replacements):
    def repl_para(p):
        if not p.runs:
            return
        txt = "".join(r.text for r in p.runs)
        novo = txt
        for a, b in replacements.items():
            novo = novo.replace(a, b)
        # Limpeza de sobras do modelo-base após substituições parciais.
        novo = novo.replace("Cristian Marcelo Schibelsky Matioli", "Cristian Marcelo Schibelsky")
        novo = novo.replace("CRISTIAN MARCELO SCHIBELSKY MATIOLI", "CRISTIAN MARCELO SCHIBELSKY")
        novo = novo.replace("cargo auxiliar administrativo", f"cargo {replacements.get('Diretor de Sub Divisão', 'Diretor de Sub Divisão')}")
        novo = novo.replace("cargo Auxiliar Administrativo", f"cargo {replacements.get('Diretor de Sub Divisão', 'Diretor de Sub Divisão')}")
        novo = novo.replace("cargo AUXILIAR ADMINISTRATIVO", f"cargo {replacements.get('DIRETOR DE SUB DIVISÃO', 'DIRETOR DE SUB DIVISÃO')}")
        novo = corrigir_gramatica_artistica_texto(novo)
        if novo != txt:
            p.runs[0].text = novo
            for r in p.runs[1:]:
                r.text = ""
    for p in doc.paragraphs:
        repl_para(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    repl_para(p)


def contexto_replacements(dados, processo_id):
    dados = enriquecer(dados, {})
    artista = dados.get("artista") or "artista indicada"
    evento = dados.get("evento") or "evento municipal"
    objeto = objeto_padrao(dados)
    valor = dados.get("valor_estimado") or ""
    processo = numero_processo(processo_id, dados)
    data_doc = data_documento_final(dados)
    secretario = dados.get("secretario") or ""
    dados = aplicar_cadastro_agente(dados)
    agente = dados.get("agente") or ""
    agente_matricula = dados.get("agente_matricula") or ""
    agente_cargo = dados.get("agente_cargo") or ""
    gestor = dados.get("gestor") or GESTOR_PADRAO
    fiscal = dados.get("fiscal") or FISCAL_PADRAO
    gestor_cargo = dados.get("gestor_cargo") or GESTOR_CARGO_PADRAO
    fiscal_cargo = dados.get("fiscal_cargo") or FISCAL_CARGO_PADRAO
    gestor_comp = dados.get("gestor_competencias") or GESTOR_COMPETENCIAS_PADRAO
    fiscal_comp = dados.get("fiscal_competencias") or FISCAL_COMPETENCIAS_PADRAO

    return {
        "21674/2026": processo,
        "Gabriela Rocha": artista,
        "GABRIELA ROCHA": artista.upper(),
        "cantora Gabriela Rocha": f"cantora {artista}",
        "Natal para Jesus": evento,
        "NATAL PARA JESUS": evento.upper(),
        "Contratação de serviços artísticos para a realização de apresentação ao vivo da cantora Gabriela Rocha, no evento “Natal para Jesus”, a ser realizado no Município de Sumaré/SP": objeto,
        "R$ 300.000,00": valor,
        "R$300.000,00": valor,
        "15 de Maio de 2026": data_doc,
        "16 de Junho de 2026": data_doc,
        "15 de Maio de 2025": data_doc,
        "Dayara Cristina Marques Matioli": agente,
        "DAYARA CRISTINA MARQUES MATIOLI": agente.upper(),
        "Dayara Cristina Marques": agente,
        "DAYARA CRISTINA MARQUES": agente.upper(),
        "Cristian Marcelo Schibelsky Matioli": agente,
        "CRISTIAN MARCELO SCHIBELSKY MATIOLI": agente.upper(),
        "Cristian Marcelo Schibelsky": agente,
        "CRISTIAN MARCELO SCHIBELSKY": agente.upper(),
        "17802": agente_matricula,
        "22019": agente_matricula,
        "Diretor de Sub Divisão": agente_cargo,
        "DIRETOR DE SUB DIVISÃO": agente_cargo.upper(),
        "cargo auxiliar administrativo": f"cargo {agente_cargo}",
        "cargo Auxiliar Administrativo": f"cargo {agente_cargo}",
        "cargo AUXILIAR ADMINISTRATIVO": f"cargo {agente_cargo.upper()}",
        "cargo auxiliar administrativo,": f"cargo {agente_cargo},",
        "cargo Auxiliar Administrativo,": f"cargo {agente_cargo},",
        "cargo AUXILIAR ADMINISTRATIVO,": f"cargo {agente_cargo.upper()},",
        "CECILIA SOUSA TEIXEIRA": secretario.upper(),
        "Secretário Municipal de Cultura e Turismo": dados.get("cargo_secretario") or "Cargo da autoridade a definir",
        "Secretária Municipal de Cultura e Turismo": dados.get("cargo_secretario") or "Cargo da autoridade a definir",
        "Secretario Municipal de Cultura e Turismo": dados.get("cargo_secretario") or "Cargo da autoridade a definir",
        "Secretaria Municipal de Cultura e Turismo": dados.get("cargo_secretario") or "Cargo da autoridade a definir",
        "Carla Andressa Dourado": gestor,
        "CARLA ANDRESSA DOURADO": gestor.upper(),
        "Talita Cristiane Carvalho": fiscal,
        "TALITA CRISTIANE CARVALHO": fiscal.upper(),
        "Sra. Cristian Marcelo Schibelsky": "Sr.(a) Cristian Marcelo Schibelsky",
        "Sra. CRISTIAN MARCELO SCHIBELSKY": "Sr.(a) CRISTIAN MARCELO SCHIBELSKY",
        "Diretora de Divisão": gestor_cargo,
        "DIRETORA DE DIVISÃO": gestor_cargo.upper(),
        "Diretora de Área": fiscal_cargo,
        "DIRETORA DE ÁREA": fiscal_cargo.upper(),
        "Vasto conhecimento sobre as rotinas dos eventos da Secretaria Municipal de Cultura": gestor_comp,
    }



def remover_paragrafo(p):
    el = p._element
    el.getparent().remove(el)


def corrigir_layout_designacao_gestor_fiscal(doc):
    """
    Ajuste específico do documento 'Dispõe sobre designação de gestor e fiscal'.
    O modelo original contém vários parágrafos vazios entre a assinatura da primeira página
    e o bloco de legislação, empurrando a palavra 'Legislação' para o rodapé da página 2.
    A correção remove esses vazios sem alterar marca d'água, cabeçalho, rodapé ou conteúdo legal.
    """
    # Remove sequência de parágrafos vazios imediatamente anterior a 'Legislação:'
    idx_leg = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().lower().startswith('legislação') or p.text.strip().lower().startswith('legislacao'):
            idx_leg = i
            break
    if idx_leg is not None:
        # remove vazios consecutivos antes da legislação, preservando no máximo uma linha de respiro
        to_remove = []
        j = idx_leg - 1
        empty_count = 0
        while j >= 0 and not doc.paragraphs[j].text.strip():
            empty_count += 1
            if empty_count > 1:
                to_remove.append(doc.paragraphs[j])
            j -= 1
        for p in to_remove:
            remover_paragrafo(p)

    # Ajusta o bloco de legislação para começar com espaçamento institucional, sem salto excessivo.
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.lower().startswith('legislação') or t.lower().startswith('legislacao'):
            # A legislação deve iniciar uma nova página de forma limpa,
            # evitando ficar isolada no rodapé da primeira página.
            p.paragraph_format.page_break_before = True
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            for r in p.runs:
                r.bold = True
            break

    # Normaliza parágrafos vazios restantes para altura mínima, evitando páginas quase em branco.
    for p in doc.paragraphs:
        if not p.text.strip():
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1

    # Compacta moderadamente o texto legal para reduzir quebra ruim sem prejudicar legibilidade.
    in_leg = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.lower().startswith('legislação') or t.lower().startswith('legislacao'):
            in_leg = True
        if in_leg:
            # Remove espaçamentos artificiais herdados do modelo (ex.: "por      1 (um)").
            texto_original = p.text
            texto_limpo = re.sub(r"[ \u00a0]{2,}", " ", texto_original).strip()
            # Remove repetição de hyperlink herdada do modelo original.
            texto_limpo = texto_limpo.replace("atribuição.art. 7º desta Lei", "atribuição.")
            # Reescreve o parágrafo quando houver espaços artificiais ou elementos hyperlink duplicados.
            if texto_limpo and texto_limpo != texto_original:
                p.clear()
                p.add_run(texto_limpo)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.0
            if t.startswith('Art.') or t.startswith('§') or t[:2] in {'I ', 'II', 'III', 'IV', 'V ', 'VI', 'XI', 'X-'}:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for r in p.runs:
                r.font.name = FONTE_PADRAO
                r._element.rPr.rFonts.set(qn('w:eastAsia'), FONTE_PADRAO)
                r.font.size = Pt(10)


def gerar_doc_auxiliar(modelo, saida, processo_id, dados, pasta_saida, extraidos=None):
    pasta_saida.mkdir(parents=True, exist_ok=True)
    doc = copiar_modelo(modelo)
    dados_aux = enriquecer(dados, extraidos or {})
    replace_doc_text(doc, contexto_replacements(dados_aux, processo_id))
    arquivo = pasta_saida / saida
    doc.save(arquivo)
    return arquivo


def gerar_designacao_etp(processo_id, dados, classificacao, pasta_saida: Path, extraidos=None):
    return gerar_doc_auxiliar("Designacao_Agente_ETP_2025.docx", "05_Designacao_Agente_ETP.docx", processo_id, dados, pasta_saida, extraidos)


def gerar_designacao_tr(processo_id, dados, classificacao, pasta_saida: Path, extraidos=None):
    return gerar_doc_auxiliar("Designacao_Agente_TR_2025.docx", "06_Designacao_Agente_TR.docx", processo_id, dados, pasta_saida, extraidos)


def aplicar_dados_gestor_fiscal_padrao(doc, dados):
    """
    Garante que o documento de designação de gestor/fiscal traga os dados definidos pelo usuário,
    sem depender de extração frágil do modelo.
    """
    repl = {
        "Nome: Carla Andressa Dourado": f"Nome: {GESTOR_PADRAO}",
        "Cargo: Diretora de Divisão": f"Cargo: {GESTOR_CARGO_PADRAO}",
        "Competências: Vasto conhecimento sobre as rotinas dos eventos da Secretaria Municipal de \nCultura": f"Competências: {GESTOR_COMPETENCIAS_PADRAO}",
        "Competências: Vasto conhecimento sobre as rotinas dos eventos da Secretaria Municipal de Cultura": f"Competências: {GESTOR_COMPETENCIAS_PADRAO}",
        "Nome: Talita Cristiane Carvalho": f"Nome: {FISCAL_PADRAO}",
        "Cargo: Diretora de Área": f"Cargo: {FISCAL_CARGO_PADRAO}",
    }
    replace_doc_text(doc, repl)

    texto = "\n".join(p.text for p in doc.paragraphs)
    if "Gestor de Contrato:" not in texto or "Fiscal do Contrato:" not in texto:
        doc.add_page_break()
        add_bold(doc, "GESTOR E FISCAL DO CONTRATO")
        add_p(doc, "Gestor de Contrato:")
        add_p(doc, f"Nome: {GESTOR_PADRAO}")
        add_p(doc, f"Cargo: {GESTOR_CARGO_PADRAO}")
        add_p(doc, f"Competências: {GESTOR_COMPETENCIAS_PADRAO}")
        add_p(doc, "Ciente, de acordo. ___________________________________________________________")
        add_p(doc, "")
        add_p(doc, "Fiscal do Contrato:")
        add_p(doc, f"Nome: {FISCAL_PADRAO}")
        add_p(doc, f"Cargo: {FISCAL_CARGO_PADRAO}")
        add_p(doc, f"Competências: {FISCAL_COMPETENCIAS_PADRAO}")
        add_p(doc, "Ciente, de acordo. _____________________________________________________")


def gerar_designacao_gestor_fiscal(processo_id, dados, classificacao, pasta_saida: Path, extraidos=None):
    pasta_saida.mkdir(parents=True, exist_ok=True)
    doc = copiar_modelo("Designacao_Gestor_Fiscal_2025.docx")
    dados_aux = enriquecer(dados, extraidos or {})
    replace_doc_text(doc, contexto_replacements(dados_aux, processo_id))
    aplicar_dados_gestor_fiscal_padrao(doc, dados_aux)
    corrigir_layout_designacao_gestor_fiscal(doc)
    arquivo = pasta_saida / "07_Designacao_Gestor_Fiscal.docx"
    doc.save(arquivo)
    return arquivo


def gerar_disponibilidade_orcamentaria(processo_id, dados, classificacao, pasta_saida: Path, extraidos=None):
    return gerar_doc_auxiliar("Disponibilidade_Orcamentaria_2025.docx", "08_Disponibilidade_Orcamentaria.docx", processo_id, dados, pasta_saida, extraidos)


def gerar_certificacao_tr(processo_id, dados, classificacao, pasta_saida: Path, extraidos=None):
    return gerar_doc_auxiliar("Certificacao_TR_Padronizado_2025.docx", "09_Certificacao_TR_Padronizado.docx", processo_id, dados, pasta_saida, extraidos)



# =========================
# LISTA DE VERIFICAÇÃO - ADESÃO À ATA
# =========================

ITENS_AGU_ADESAO_ATA = [
    "Houve abertura de processo administrativo?",
    "Foi adotada a forma eletrônica para o processo administrativo ou, caso adotada forma em papel, houve a devida justificativa?",
    "Consta documento de formalização de demanda?",
    "Foi certificado que objeto da contratação está contemplado no Plano de Contratações Anual?",
    "Foi certificado que objeto da contratação está compatível com as leis orçamentárias?",
    "Há Estudo Técnico Preliminar?",
    "O estudo técnico preliminar contém as informações que bem caracterizam a contratação, tais como o quantitativo demandado e o local de entrega do bem ou de prestação do serviço?",
    "Foi apresentada justificativa da vantagem da adesão, conforme art. 31, I, do Decreto nº 11.462/2023?",
    "Os valores registrados estão compatíveis com os valores praticados pelo mercado?",
    "O fornecedor aceitou o pedido de adesão?",
    "Houve aceitação da adesão pelo órgão ou entidade gerenciadora?",
    "A ata a que se pretende aderir é gerenciada por órgão ou entidade da Administração Pública federal?",
    "Foi observado o limite de 50% dos quantitativos registrados na ata de registro de preços para o órgão gerenciador e para os órgãos participantes?",
    "A adesão será formalizada dentro do prazo de 90 dias, contado da autorização do órgão ou da entidade gerenciadora, observado o prazo de vigência da ata?",
    "A contratação será formalizada por instrumento contratual, emissão de nota de empenho de despesa, autorização de compra ou outro instrumento hábil?",
    "O instrumento que será adotado será firmado dentro do prazo de validade da ata de registro de preços?",
    "Foram consultados SICAF, CEIS, CNJ e TCU, juntando-se os respectivos comprovantes relacionados ao fornecedor?",
    "Consta dos autos consulta ao CADIN?",
    "Foi consultado o Guia Nacional de Contratações Sustentáveis da Consultoria-Geral da União para inserção dos critérios de sustentabilidade?",
]


def eh_adesao_ata(classificacao, dados=None):
    texto = " ".join([
        str((classificacao or {}).get("tipo", "")),
        str((classificacao or {}).get("modalidade", "")),
        str((classificacao or {}).get("categoria", "")),
        str((dados or {}).get("descricao", "")),
    ]).lower()
    return "adesão à ata" in texto or "adesao" in texto and "ata" in texto or "registro de preços" in texto or "registro de precos" in texto


def gerar_lista_verificacao_adesao_ata(processo_id, dados, classificacao, pasta_saida: Path, extraidos=None):
    dados = enriquecer(dados, extraidos)
    pasta_saida.mkdir(parents=True, exist_ok=True)
    arquivo = pasta_saida / "10_Lista_Verificacao_Adesao_Ata_AGU.docx"

    doc = Document()
    p = doc.add_paragraph("[ÓRGÃO OU ENTIDADE PÚBLICA]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
        run.font.name = FONTE_PADRAO
        run.font.size = Pt(11)

    add_bold(doc, "LISTA DE VERIFICAÇÃO")
    add_p(doc, "(Adesão à ata de registro de preços – Lei nº 14.133/2021)")
    add_p(doc, "Modelo incorporado como referência com base na Lista de Verificação de Adesão à Ata de Registro de Preços da AGU/CNMLC, atualização SET/2024.")
    add_p(doc, f"Processo Administrativo n° {numero_processo(processo_id, dados)}")
    add_p(doc, f"Secretaria: {dados.get('secretaria') or 'Secretaria requisitante'}")
    add_p(doc, "")

    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = "Table Grid"
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tabela.rows[0].cells
    hdr[0].text = "Nº"
    hdr[1].text = "Aspectos gerais pertinentes à adesão à ata de registro de preços"
    hdr[2].text = "Atende plenamente a exigência?"
    hdr[3].text = "Indicação do local do processo em que foi atendida a exigência (doc. / fls. / sistema)"

    for i, item in enumerate(ITENS_AGU_ADESAO_ATA, start=1):
        row = tabela.add_row().cells
        row[0].text = str(i)
        row[1].text = item
        row[2].text = "Resposta"
        row[3].text = ""

    for row in tabela.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = FONTE_PADRAO
                    r.font.size = Pt(9)

    add_p(doc, "")
    add_p(doc, "Observação: a coluna de resposta deve ser preenchida com Sim, Não ou Não se aplica, conforme análise do caso concreto.")
    add_p(doc, "Esta lista deve ser juntada aos autos antes da remessa ao órgão de assessoramento jurídico, sem prejuízo das exigências municipais específicas.")
    add_p(doc, "")
    add_p(doc, f"Sumaré, {val(data_documento_final(dados))}.", style="assinatura")
    add_p(doc, "__________________________________", style="assinatura")
    resp = clean(dados.get("responsavel_demanda") or dados.get("secretario") or "")
    cargo = clean(dados.get("cargo_secretario") or dados.get("secretaria") or "")
    add_p(doc, resp.upper() if resp else "RESPONSÁVEL PELA INSTRUÇÃO", style="assinatura")
    add_p(doc, cargo.upper() if cargo else "CARGO/FUNÇÃO", style="assinatura")

    aplicar_layout_documento(doc, "LISTA DE VERIFICAÇÃO")
    doc.save(arquivo)
    return arquivo


def gerar_processo_completo(processo_id, dados, classificacao, pasta_saida: Path, extraidos=None):
    pasta_saida.mkdir(parents=True, exist_ok=True)
    arquivos = [
        gerar_dfd(processo_id, dados, classificacao, pasta_saida, extraidos),
        gerar_designacao_etp(processo_id, dados, classificacao, pasta_saida, extraidos),
        gerar_etp(processo_id, dados, classificacao, pasta_saida, extraidos),
        gerar_designacao_tr(processo_id, dados, classificacao, pasta_saida, extraidos),
        gerar_tr(processo_id, dados, classificacao, pasta_saida, extraidos),
        gerar_certificacao_tr(processo_id, dados, classificacao, pasta_saida, extraidos),
        gerar_designacao_gestor_fiscal(processo_id, dados, classificacao, pasta_saida, extraidos),
        gerar_disponibilidade_orcamentaria(processo_id, dados, classificacao, pasta_saida, extraidos),
    ]
    if eh_adesao_ata(classificacao, dados):
        arquivos.append(gerar_lista_verificacao_adesao_ata(processo_id, dados, classificacao, pasta_saida, extraidos))
    return arquivos