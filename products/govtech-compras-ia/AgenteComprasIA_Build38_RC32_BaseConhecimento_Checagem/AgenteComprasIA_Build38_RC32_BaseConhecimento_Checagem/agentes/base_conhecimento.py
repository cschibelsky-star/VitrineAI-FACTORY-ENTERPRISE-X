from pathlib import Path
import json
import re
import unicodedata

try:
    from docx import Document
except Exception:
    Document = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None


ROOT = Path(__file__).resolve().parents[1]
BASE_DIR = ROOT / "base_conhecimento"
MODELOS_DIR = ROOT / "modelos"
MODELOS_PREFEITURA_DIR = ROOT / "modelos_prefeitura"


def normalizar(txt):
    txt = str(txt or "").lower()
    txt = unicodedata.normalize("NFKD", txt)
    txt = "".join(c for c in txt if not unicodedata.combining(c))
    txt = re.sub(r"[^a-z0-9]+", " ", txt)
    return re.sub(r"\s+", " ", txt).strip()


def ler_docx(path):
    if Document is None:
        return ""
    try:
        doc = Document(str(path))
        partes = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        partes.append(cell.text)
        return "\n".join(partes)
    except Exception:
        return ""


def ler_pdf(path, max_paginas=20):
    if PdfReader is None:
        return ""
    try:
        reader = PdfReader(str(path))
        partes = []
        for page in reader.pages[:max_paginas]:
            try:
                partes.append(page.extract_text() or "")
            except Exception:
                pass
        return "\n".join(partes)
    except Exception:
        return ""


def ler_texto_arquivo(path):
    path = Path(path)
    suf = path.suffix.lower()
    if suf in [".txt", ".md", ".json", ".jsonl", ".csv"]:
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""
    if suf == ".docx":
        return ler_docx(path)
    if suf == ".pdf":
        return ler_pdf(path)
    return ""


def carregar_referencias():
    """
    Carrega referências internas:
    - base_conhecimento/homologada;
    - modelos/adesao_ata;
    - modelos_prefeitura.
    Não envia nada para fora; apenas monta um índice local simples.
    """
    refs = []
    pastas = [
        BASE_DIR,
        MODELOS_DIR,
        MODELOS_PREFEITURA_DIR,
    ]

    for pasta in pastas:
        if not pasta.exists():
            continue
        for arq in pasta.rglob("*"):
            if not arq.is_file():
                continue
            if arq.name.startswith("~$"):
                continue
            if arq.suffix.lower() not in [".docx", ".pdf", ".txt", ".md", ".json", ".jsonl"]:
                continue
            texto = ler_texto_arquivo(arq)
            if not texto and arq.suffix.lower() not in [".jsonl", ".json"]:
                continue
            rel = str(arq.relative_to(ROOT))
            refs.append({
                "arquivo": rel,
                "nome": arq.name,
                "tipo": "modelo_adesao_ata" if "adesao_ata" in rel.lower() else "base_homologada" if "homologada" in rel.lower() else "modelo_prefeitura",
                "texto": texto[:25000],
                "texto_norm": normalizar(texto + " " + rel),
            })
    return refs


def buscar_referencias(consulta, limite=5, filtro_tipo=None):
    refs = carregar_referencias()
    q_norm = normalizar(consulta)
    termos = [t for t in q_norm.split() if len(t) >= 4]
    resultados = []

    for ref in refs:
        if filtro_tipo and ref.get("tipo") != filtro_tipo:
            continue
        texto = ref.get("texto_norm", "")
        score = 0
        for t in termos:
            if t in texto:
                score += 3
        # bônus por nomes/tipos muito relevantes
        if "adesao" in q_norm and "ata" in q_norm and "adesao" in texto and "ata" in texto:
            score += 25
        if "registro de precos" in q_norm and "registro de precos" in texto:
            score += 20
        if "lei 14 133" in q_norm and ("14 133" in texto or "14133" in texto):
            score += 10
        if score > 0:
            r = dict(ref)
            r["score"] = score
            resultados.append(r)

    resultados.sort(key=lambda x: x["score"], reverse=True)
    return resultados[:limite]


def resumo_referencias(consulta, limite=5):
    refs = buscar_referencias(consulta, limite=limite)
    if not refs:
        return "Nenhuma referência local relevante localizada na base de conhecimento."
    linhas = []
    for i, ref in enumerate(refs, start=1):
        amostra = " ".join((ref.get("texto") or "").split())[:800]
        linhas.append(
            f"{i}. {ref.get('arquivo')} | tipo={ref.get('tipo')} | score={ref.get('score')}\n"
            f"   Trecho: {amostra}"
        )
    return "\n".join(linhas)


def checklist_base_adesao_ata():
    """
    Retorna checklist consolidado com base no modelo AGU SET/2024 e complementos municipais.
    """
    return [
        "Abertura de processo administrativo",
        "Forma eletrônica ou justificativa para processo em papel",
        "Documento de Formalização de Demanda - DFD",
        "Compatibilidade com o Plano de Contratações Anual - PCA",
        "Compatibilidade com leis orçamentárias",
        "Estudo Técnico Preliminar - ETP",
        "ETP com quantitativo demandado e local de entrega/prestação",
        "Justificativa da vantagem da adesão",
        "Compatibilidade dos valores registrados com o mercado",
        "Aceite do fornecedor",
        "Aceitação/autorização do órgão gerenciador",
        "Ata gerenciada por órgão ou entidade federal, quando aplicável",
        "Limite de 50% dos quantitativos registrados",
        "Formalização em até 90 dias da autorização do gerenciador",
        "Contrato, empenho, autorização de compra ou instrumento hábil",
        "Instrumento firmado dentro da validade da ata",
        "Consultas SICAF, CEIS, CNJ e TCU",
        "Consulta ao CADIN",
        "Consulta ao Guia Nacional de Contratações Sustentáveis",
        "Edital da licitação de origem",
        "Ata de Registro de Preços",
        "Publicação da Ata e homologação",
        "TR original da licitação",
        "Proposta vencedora",
        "Parecer jurídico",
        "Autorização da autoridade competente",
        "Publicação do extrato do contrato",
    ]


def gerar_contexto_base_para_revisor(dados, classificacao=None):
    texto_busca = " ".join([
        str((dados or {}).get("descricao", "")),
        str((dados or {}).get("evento", "")),
        str((classificacao or {}).get("tipo", "")),
        str((classificacao or {}).get("modalidade", "")),
        "adesão ata registro preços lei 14.133 checklist" if "ata" in normalizar((classificacao or {}).get("tipo", "")) else "",
    ])
    return resumo_referencias(texto_busca, limite=4)