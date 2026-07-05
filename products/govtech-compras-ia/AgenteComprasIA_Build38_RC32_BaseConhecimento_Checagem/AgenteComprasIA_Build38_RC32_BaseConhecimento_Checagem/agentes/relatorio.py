from pathlib import Path
from datetime import datetime
from docx import Document


def gerar_relatorio_pendencias(resultados, resumo, extraidos, classificacao, pasta_saida: Path):
    pasta_saida.mkdir(parents=True, exist_ok=True)
    arquivo = pasta_saida / "Relatorio_Conferencia_Documental.docx"
    doc = Document()

    doc.add_heading("RELATÓRIO TÉCNICO DE CONFERÊNCIA DOCUMENTAL", level=1)
    doc.add_paragraph("MUNICÍPIO DE SUMARÉ")
    doc.add_paragraph("Secretaria Municipal de Cultura e Turismo")
    doc.add_paragraph(f"Data da conferência: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    doc.add_heading("1. Síntese da conferência", level=2)
    doc.add_paragraph(f"Completude documental externa: {resumo.get('percentual')}%")
    doc.add_paragraph(f"Documentos externos identificados: {resumo.get('encontrados')}")
    doc.add_paragraph(f"Pendências externas: {resumo.get('pendentes')}")
    doc.add_paragraph(f"Pendências críticas externas: {resumo.get('criticos')}")
    doc.add_paragraph(f"Situação preliminar: {resumo.get('semaforo')}")

    doc.add_heading("2. Dados extraídos do processo", level=2)
    for k, v in (extraidos or {}).items():
        if k not in ("valores_localizados", "evidencias"):
            doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("3. Documentos externos encontrados", level=2)
    encontrou = False
    for r in resultados:
        if r["status"] == "Encontrado":
            encontrou = True
            doc.add_paragraph(f"✓ {r['item']} — {r.get('arquivo','')}")
    if not encontrou:
        doc.add_paragraph("Nenhum documento externo identificado.")

    doc.add_heading("4. Pendências externas", level=2)
    tem = False
    for r in resultados:
        if r["status"] != "Encontrado":
            tem = True
            marca = "CRÍTICO" if r["criticidade"] == "critico" else "ATENÇÃO"
            doc.add_paragraph(f"{marca}: {r['item']}")
    if not tem:
        doc.add_paragraph("Não foram identificadas pendências externas relevantes no checklist aplicado.")

    doc.add_heading("5. Documentos técnicos gerados pelo sistema", level=2)
    doc.add_paragraph("Os documentos abaixo não constituem pendências documentais externas. Eles integram a fase de instrução técnica e podem ser gerados a partir dos dados e evidências do processo:")
    for item, _ in classificacao.get("documentos_gerados_sistema", []):
        doc.add_paragraph(f"• {item}")

    doc.add_heading("6. Observação técnica", level=2)
    doc.add_paragraph(
        "A conferência documental não substitui a análise da unidade requisitante, do setor de compras, do controle interno ou da assessoria jurídica. "
        "O relatório indica a presença ou ausência de documentos relevantes para instrução processual e subsidia a geração de DFD, ETP e TR."
    )

    doc.save(arquivo)
    return arquivo